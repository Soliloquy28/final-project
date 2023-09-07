from __future__ import print_function
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import os
import json
import datetime
import math
from screeninfo import get_monitors
from docx import Document
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
from RailDetector import *
import Global


# Metrics Selection Page frame class
class MetricsSelectionPage(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        self.back_button = tk.Button(self, text="Back", command=self.back_to_process_page, cursor="hand2")
        self.back_button.place(relx=0.05, rely=0.05, anchor='nw')
        self.proximity_label = tk.Label(self, text="Proximity Analysis:")
        self.proximity_label.place(relx=0.5, rely=0.25, anchor='center')
        self.proximity_btn = tk.Button(self, text="Prox & CSEM & Cs", command=self.prox_csem_cs)
        self.proximity_btn.place(relx=0.5, rely=0.35, anchor='center')
        self.coordination_label = tk.Label(self, text="Coordination Analysis:")
        self.coordination_label.place(relx=0.5, rely=0.5, anchor='center')
        self.coordination_btn = tk.Button(self, text="Correlation & DI & Cs", command=self.rv_di_cs)
        self.coordination_btn.place(relx=0.5, rely=0.6, anchor='center')
        # Save sorting data index for the tree
        self.tree_dict = dict()

    def back_to_process_page(self):
        from ProcessingPage import ProcessingPage
        self.master.show_frame(ProcessingPage)


    def prox_csem_cs(self):
        if any(os.scandir(self.master.processed_video_path_dynamic)):

            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            tracking_method = data['trackingmethod']
            video_path = data["processedvideo"]
            # distance_list = data["distancelist"]
            object_1_coordinate = data['object1coordinatelist']
            object_2_coordinate = data['object2coordinatelist']
            # distance_list.sort()
            min_distance = min(data["distancelist"])
            max_distance = max(data["distancelist"])
            if tracking_method == "Multiple object tracking":
                width_ratio = Global.multiple_width_ratio
                height_ratio = Global.multiple_height_ratio
            elif tracking_method == "Blob tracking":
                width_ratio = Global.blob_width_ratio
                height_ratio = Global.blob_height_ratio
            all_distance_list = []
            for cor1 in object_1_coordinate:
                for cor2 in object_2_coordinate:
                    distance = round(math.sqrt(((cor1[0] - cor2[0]) / width_ratio) ** 2 + ((cor1[1] - cor2[1]) / height_ratio) ** 2), 2)
                    all_distance_list.append(distance)
            cap = cv.VideoCapture(video_path)
            # Check if the video was successfully opened
            if not cap.isOpened():
                exit()
            # Get the total number of frames in the video
            total_frame_count = len(data["distancelist"])
            self.master.total_frame = total_frame_count
            self.screen_width = get_monitors()[0].width
            self.screen_height = get_monitors()[0].height
            self.prox_csem_cs_toplevel = tk.Toplevel(self.master)
            self.prox_csem_cs_toplevel.geometry(f"{int(self.screen_width)}x{int(self.screen_height - 200)}")
            # Set the width for each frame
            frame_width = int((self.screen_width) * 3 / 8)
            frame_height = int(self.screen_height - 200)
            # Create the left frame
            self.frame_1 = tk.Frame(self.prox_csem_cs_toplevel, width=frame_width)
            self.frame_1.pack(side='left', fill='both', expand=True)
            self.metrics_canvas_1 = tk.Canvas(self.frame_1, width=int((self.screen_width) * 3 / 8) - 25)
            self.metrics_scrollbar_1 = ttk.Scrollbar(self.frame_1, orient=tk.VERTICAL, command=self.metrics_canvas_1.yview)
            self.canvas_1_frame = tk.Frame(self.metrics_canvas_1)
            self.canvas_1_frame.bind("<Configure>", lambda e: self.metrics_canvas_1.configure(scrollregion=self.metrics_canvas_1.bbox("all")))
            self.metrics_canvas_1.create_window((int(((self.screen_width * 3 / 8) - 25) / 2), 0), window=self.canvas_1_frame, anchor="center")
            self.metrics_canvas_1.configure(yscrollcommand=self.metrics_scrollbar_1.set)
            self.metrics_scrollbar_1.pack(side=tk.RIGHT, fill=tk.Y)
            self.metrics_canvas_1.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            self.pro_label_image_frame = tk.Frame(self.canvas_1_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.pro_label_image_frame.pack()
            self.pro_label = tk.Label(self.pro_label_image_frame, text="Proximity index (Prox)")
            self.pro_label.place(relx=0.5, rely=0.15, anchor="center")
            photo_1 = Global.photo_place("amProx.png", int(frame_width * 4.3 / 10), int(frame_width * 4.3 / 20))
            self.image_label_1 = tk.Label(self.pro_label_image_frame, image=photo_1)
            # Maintain a reference to the image
            self.image_label_1.image = photo_1
            # Place the image just below self.pro_label
            self.image_label_1.place(relx=0.5, rely=0.35, anchor="n")
            self.pro_threshold_graph_frame = tk.Frame(self.canvas_1_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 2))
            self.pro_threshold_graph_frame.pack()
            self.pro_distance_threshold_graph_frame = tk.Frame(self.pro_threshold_graph_frame)
            self.pro_distance_threshold_graph_frame.place(relx=0.5, rely=0.5, anchor='center')
            self.pro_distance_threshold_graph = Figure(figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.pro_subplot = self.pro_distance_threshold_graph.add_subplot(111)
            self.pro_distance_list = data["distancelist"]
            start_distance = int(min(self.pro_distance_list))
            final_distance = int(max(self.pro_distance_list))
            self.total_frame = data["totalframe"]
            start_threshold = start_distance - 3
            final_threshold = final_distance + 3
            # The sequence starts at start_threshold
            # Each subsequent number increases by the step
            # size until but not including final_threshold
            step = 0.5
            # Create a sequence using numpy.arange
            self.threshold_sequence = np.arange(start_threshold, final_threshold, step)
            prox_value_list = []
            for thre in self.threshold_sequence:
                kernel = 0
                for dis in data["distancelist"]:
                    if dis < thre:
                        kernel = kernel + 1
                prox_value = kernel / self.total_frame
                prox_value_list.append(prox_value)
            self.pro_x_cor = self.threshold_sequence
            self.pro_y_cor = prox_value_list
            self.pro_subplot.plot(self.pro_x_cor, self.pro_y_cor, color='r')
            self.pro_subplot.set_xlabel("Threshold (cm)")
            self.pro_subplot.set_ylabel("Proximity Index")
            self.pro_distance_threshold_graph.tight_layout()
            self.pro_distance_threshold_graph.subplots_adjust(top=0.9)
            self.canvas_prox = FigureCanvasTkAgg(self.pro_distance_threshold_graph, master=self.pro_threshold_graph_frame)  # A tk.DrawingArea.
            self.canvas_prox.draw()
            self.canvas_prox.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
            self.prox_line = None
            self.prox_annotation = None
            self.canvas_prox.mpl_connect('button_press_event', self.draw_line_1)
            self.canvas_prox.mpl_connect('motion_notify_event', self.move_cursor_1)
            self.pro_label_entry_frame = tk.Frame(self.canvas_1_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.pro_label_entry_frame.pack()
            # pro T value
            self.pro_T_label = tk.Label(self.pro_label_entry_frame, text=f"T = {self.master.total_frame}")
            self.pro_T_label.place(relx=0.5, rely=0.2, anchor="center")
            # pro distance
            self.pro_distance_label = tk.Label(self.pro_label_entry_frame, text=f"The distance ranges from {min_distance}cm to {max_distance}cm.")
            self.pro_distance_label.place(relx=0.5, rely=0.4, anchor="center")
            # threshold entry frame
            self.pro_threshold_frame = tk.Frame(self.pro_label_entry_frame, width=frame_width * 7 / 10)
            self.pro_threshold_frame.place(relx=0.5, rely=0.6, anchor="center")
            # δ entry
            self.pro_threshold_label = tk.Label(self.pro_threshold_frame, text="Threshold (δ): ")
            self.pro_threshold_label.pack(side=tk.LEFT)
            self.pro_threshold_entry = tk.Entry(self.pro_threshold_frame, textvariable=self.master.proximity_threshold)
            self.pro_threshold_entry.pack(side=tk.RIGHT)
            # δ submit
            self.pro_submit = tk.Button(self.pro_label_entry_frame, text="Submit", command=self.submit_pro_threshold)
            self.pro_submit.place(relx=0.5, rely=0.8, anchor="center")
            self.pro_tree_delete_frame = tk.Frame(self.canvas_1_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.pro_tree_delete_frame.pack()
            # proximity treeview and scrollbar frame
            self.pro_tree_scroll_frame = tk.Frame(self.pro_tree_delete_frame)
            self.pro_tree_scroll_frame.place(relx=0.5, rely=0.4, relwidth=0.9, relheight=0.7, anchor="center")
            self.pro_tree = ttk.Treeview(self.pro_tree_scroll_frame, columns=("c1", "c2", "c3", "c4"), show='headings', selectmode="extended")
            self.pro_tree.column("c1", anchor=tk.CENTER, width=int(frame_width * 1 / 6))
            self.pro_tree.heading("c1", text="Author", command=lambda _col="c1": self.sortby(_col, 0, self.pro_tree))
            self.pro_tree.column("c2", anchor=tk.CENTER, width=int(frame_width * 1 / 6))
            self.pro_tree.heading("c2", text="δ", command=lambda _col="c2": self.sortby(_col, 0, self.pro_tree))
            self.pro_tree.column("c3", anchor=tk.CENTER, width=int(frame_width * 1 / 6))
            self.pro_tree.heading("c3", text="Prox", command=lambda _col="c3": self.sortby(_col, 0, self.pro_tree))
            self.pro_tree.column("c4", width=0, minwidth=0, stretch='no')
            self.pro_tree.heading("c4", text="Timestamp", command=lambda _col="c4": self.sortby(_col, 0, self.pro_tree))
            self.pro_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            # Create a vertical scrollbar
            self.pro_tree_vsb = ttk.Scrollbar(self.pro_tree_scroll_frame, orient="vertical",command=self.pro_tree.yview)
            self.pro_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.pro_tree.configure(yscrollcommand=self.pro_tree_vsb.set)
            # delete item in treeview
            self.pro_tree_delete_btn = tk.Button(self.pro_tree_delete_frame, text="Delete", command=self.delete_pro_item)
            self.pro_tree_delete_btn.place(relx=0.5, rely=0.9, anchor="center")
            # prox pre-insert
            for prox in data["prox"]:
                self.pro_tree.insert('', 'end', values=(prox["author"], prox["δ"], prox["prox"], prox["timestamp"]))

            # Create the middle frame
            self.frame_2 = tk.Frame(self.prox_csem_cs_toplevel, width=frame_width)
            self.frame_2.pack(side='left', fill='both', expand=True)
            self.metrics_canvas_2 = tk.Canvas(self.frame_2, width=int((self.screen_width) * 3 / 8) - 25)
            self.metrics_scrollbar_2 = ttk.Scrollbar(self.frame_2, orient=tk.VERTICAL,
                                                     command=self.metrics_canvas_2.yview)
            self.canvas_2_frame = tk.Frame(self.metrics_canvas_2)
            self.canvas_2_frame.bind("<Configure>", lambda e: self.metrics_canvas_2.configure(
                scrollregion=self.metrics_canvas_2.bbox("all")))
            self.metrics_canvas_2.create_window((int(((self.screen_width * 3 / 8) - 25) / 2), 0),
                                                window=self.canvas_2_frame, anchor="center")
            self.metrics_canvas_2.configure(yscrollcommand=self.metrics_scrollbar_2.set)
            self.metrics_scrollbar_2.pack(side=tk.RIGHT, fill=tk.Y)
            self.metrics_canvas_2.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            self.csem_label_image_frame = tk.Frame(self.canvas_2_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.csem_label_image_frame.pack()
            self.csem_label = tk.Label(self.csem_label_image_frame, text="Standardized cross sampled entropy (CSEM)")
            self.csem_label.place(relx=0.5, rely=0.15, anchor='center')
            photo_2 = Global.photo_place("amCSEM.png", int(frame_width * 6.3 / 10), int(frame_width * 6.3 / 20))
            self.image_label_2 = tk.Label(self.csem_label_image_frame, image=photo_2)
            self.image_label_2.image = photo_2
            self.image_label_2.place(relx=0.5, rely=0.35, anchor='n')
            self.csem_threshold_graph_frame = tk.Frame(self.canvas_2_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 2))
            self.csem_threshold_graph_frame.pack()
            self.csem_distance_list = data["distancelist"]
            self.csem_distance_threshold_graph = Figure(
                figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.csem_subplot = self.csem_distance_threshold_graph.add_subplot(111)
            csem_threshold_sequence = self.threshold_sequence


            def group_max(lst, group_size):
                max_vals = []
                for i in range(0, len(lst) - group_size + 1):
                    max_vals.append(max(lst[i:i + group_size]))
                return max_vals
            self.csem_x_cor = []
            self.csem_y_cor = []
            for threshold in csem_threshold_sequence:
                max_m_list = []
                # 1 - self.total_frame - 1
                # 0 - self.total_frame - 2
                for m in range(self.total_frame - 1):
                    m = m + 1
                    Nm = 0
                    max_list = group_max(self.csem_distance_list, m)
                    for max1 in max_list:
                        if max1 < threshold:
                            Nm = Nm + 1
                    if Nm > 0:
                        max_m_list.append(m)
                    elif Nm == 0:
                        break
                if max_m_list != []:
                    max_m_value = max_m_list[-1]
                    csem_value = max_m_value / (self.total_frame - 1)
                else:
                    csem_value = 0
                self.csem_x_cor.append(threshold)
                self.csem_y_cor.append(csem_value)
            # If you have many lines and don't want to manually specify
            # the color for each one, Matplotlib will automatically assign
            # different colors to your lines.
            self.csem_subplot.plot(self.csem_x_cor, self.csem_y_cor, color='r')
            self.csem_subplot.set_xlabel("Threshold (cm)")
            self.csem_subplot.set_ylabel("CSEM")
            self.csem_distance_threshold_graph.tight_layout()
            self.csem_distance_threshold_graph.subplots_adjust(top=0.9)
            self.canvas_csem = FigureCanvasTkAgg(self.csem_distance_threshold_graph,
                                                 master=self.csem_threshold_graph_frame)
            self.canvas_csem.draw()
            self.canvas_csem.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
            self.csem_line = None
            self.csem_annotation = None
            self.canvas_csem.mpl_connect('button_press_event', self.draw_line_2)
            self.canvas_csem.mpl_connect('motion_notify_event', self.move_cursor_2)
            self.csem_label_entry_frame = tk.Frame(self.canvas_2_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.csem_label_entry_frame.pack()
            # csem T value
            self.csem_T_label = tk.Label(self.csem_label_entry_frame, text=f"T = {self.master.total_frame}")
            self.csem_T_label.place(relx=0.5, rely=0.2, anchor='center')
            # csem distance
            self.csem_distance_label = tk.Label(self.csem_label_entry_frame,
                                                text=f"The distance ranges from {min_distance}cm to {max_distance}cm.")
            self.csem_distance_label.place(relx=0.5, rely=0.4, anchor='center')
            # threshold entry frame
            self.csem_threshold_frame = tk.Frame(self.csem_label_entry_frame, width=frame_width * 7 / 10)
            self.csem_threshold_frame.place(relx=0.5, rely=0.6, anchor='center')
            # δ entry
            self.csem_threshold_label = tk.Label(self.csem_threshold_frame, text="Threshold (δ): ")
            self.csem_threshold_label.pack(side=tk.LEFT)
            self.csem_threshold_entry = tk.Entry(self.csem_threshold_frame, textvariable=self.master.csem_threshold)
            self.csem_threshold_entry.pack(side=tk.RIGHT)
            # δ submit
            self.csem_submit = tk.Button(self.csem_label_entry_frame, text="Submit", command=self.submit_csem_threshold)
            self.csem_submit.place(relx=0.5, rely=0.8, anchor='center')
            self.csem_tree_delete_frame = tk.Frame(self.canvas_2_frame, width=int(frame_width * 12 / 13), height=int(frame_height * 1 / 3))
            self.csem_tree_delete_frame.pack()
            # csem treeview and scrollbar frame
            self.csem_tree_scroll_frame = tk.Frame(self.csem_tree_delete_frame)
            self.csem_tree_scroll_frame.place(relx=0.5, rely=0.4, relwidth=0.9, relheight=0.7, anchor="center")
            self.csem_tree = ttk.Treeview(self.csem_tree_scroll_frame, columns=("c1", "c2", "c3", "c4", "c5"), show='headings', height=8, selectmode="extended")
            self.csem_tree.column("c1", anchor=tk.CENTER, width=int(frame_width * 1 / 10))
            self.csem_tree.heading("c1", text="Author", command=lambda _col="c1": self.sortby(_col, 0, self.csem_tree))
            self.csem_tree.column("c2", anchor=tk.CENTER, width=int(frame_width * 1 / 10))
            self.csem_tree.heading("c2", text="δ", command=lambda _col="c2": self.sortby(_col, 0, self.csem_tree))
            self.csem_tree.column("c3", anchor=tk.CENTER, width=int(frame_width * 1 / 10))
            self.csem_tree.heading("c3", text="max(m)", command=lambda _col="c3": self.sortby(_col, 0, self.csem_tree))
            self.csem_tree.column("c4", anchor=tk.CENTER, width=int(frame_width * 2 / 9))
            self.csem_tree.heading("c4", text="CSEM", command=lambda _col="c4": self.sortby(_col, 0, self.csem_tree))
            self.csem_tree.column("c5", width=0, minwidth=0, stretch='no')
            self.csem_tree.heading("c5", text="Timestamp", command=lambda _col="c5": self.sortby(_col, 0, self.csem_tree))
            self.csem_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            # Create a vertical scrollbar
            self.csem_tree_vsb = ttk.Scrollbar(self.csem_tree_scroll_frame, orient="vertical", command=self.csem_tree.yview)
            self.csem_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.csem_tree.configure(yscrollcommand=self.csem_tree_vsb.set)
            # delete item in treeview
            self.csem_tree_delete_btn = tk.Button(self.csem_tree_delete_frame, text="Delete", command=self.delete_csem_item)
            self.csem_tree_delete_btn.place(relx=0.5, rely=0.9, anchor='center')
            # csem pre-insert
            for csem in data["csem"]:
                self.csem_tree.insert('', 'end', values=(csem["author"], csem["δ"], csem["max(m)"], csem["CSEM"], csem["timestamp"]))

            # Create the right frame
            self.frame_3 = tk.Frame(self.prox_csem_cs_toplevel, width=int(frame_width * 2 / 3))
            self.frame_3.pack(side='left', fill='both', expand=True)
            self.cs_label = tk.Label(self.frame_3, text="Coefficient of Sociality (Cs)")
            self.cs_label.place(relx=0.5, rely=0.05, anchor='center')
            photo_3 = Global.photo_place("amCs.png", int(frame_width / 2), int(frame_width / 3))
            self.image_label_3 = tk.Label(self.frame_3, image=photo_3)
            self.image_label_3.image = photo_3
            self.image_label_3.place(relx=0.5, rely=0.1, anchor='n')
            # cs T value
            self.cs_T_label = tk.Label(self.frame_3, text=f"T = {self.master.total_frame}")
            self.cs_T_label.place(relx=0.5, rely=0.4, anchor='center')
            # cs DO value
            DO_value = sum(data["distancelist"]) / total_frame_count
            self.cs_do_label = tk.Label(self.frame_3, text=f"D_O = {DO_value}")
            self.cs_do_label.place(relx=0.5, rely=0.45, anchor='center')
            # cs DE value
            DE_value = sum(all_distance_list) / total_frame_count ** 2
            self.cs_de_label = tk.Label(self.frame_3, text=f"D_E = {DE_value}")
            self.cs_de_label.place(relx=0.5, rely=0.5, anchor='center')
            # cs Cs value
            Cs_value = (DE_value - DO_value) / (DE_value + DO_value)
            self.cs_cs_label = tk.Label(self.frame_3, text=f"Cs = {Cs_value}")
            self.cs_cs_label.place(relx=0.5, rely=0.55, anchor='center')
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
                # Update the data
                data['cs'] = [{
                    "D_O": DO_value,
                    "D_E": DE_value,
                    "Cs": Cs_value
                }]
            # Write the updated data back to the file
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)
            # Write to a Word document
            # Create a CS Word document and edit it to write CS results
            # Add prox to metrics file as .docx
            frame_cs_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Cs.docx'
            # os.path.exists() is a method in the os module in Python used to
            # check if a specified path (file or directory) exists
            # It returns True if the path exists and False if it does not
            if os.path.exists(frame_cs_docx_path):
                doc = Document(frame_cs_docx_path)
            else:
                # Create a Document object
                doc = Document()
            # Iterate through all paragraphs in the document and delete them
            for element in doc.element.body:
                doc.element.body.remove(element)
            doc.add_paragraph(
                # Add a paragraph to the document
                f'D_O = {DO_value}\nD_E = {DE_value}\nCs = {Cs_value}')
            # Save the document as a .docx file
            doc.save(frame_cs_docx_path)
            def move_scrollbar_to_top():
                self.metrics_canvas_1.yview_moveto(0)
                self.metrics_canvas_2.yview_moveto(0)
            # The after method will call the move_scrollbar_to_top function after 100 milliseconds.
            self.master.after(100, move_scrollbar_to_top)

        else:
            messagebox.showwarning("Warning", "Please process the video first")


    def move_cursor_1(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.pro_subplot.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_prox.draw()
        else:
            self.pro_subplot.set_title("")
            self.canvas_prox.draw()


    def draw_line_1(self, event):
        # Check if the mouse click event is within the coordinates of the image axis
        if event.inaxes is not None:
            x_value = event.xdata
            if self.prox_line:
                self.prox_line.remove()
            if self.prox_annotation:
                self.prox_annotation.remove()
            self.prox_line = self.pro_subplot.axvline(x=x_value, color='g')
            # Find the y value closest to the clicked point's x coordinate
            y_interp = np.interp(x_value, self.pro_x_cor, self.pro_y_cor)
            # Display coordinates
            self.prox_annotation = self.pro_subplot.annotate(f'({x_value:.2f}, {y_interp:.2f})', xy=(x_value, y_interp), xytext=(x_value, y_interp))
            # Redraw the canvas
            self.canvas_prox.draw()


    def move_cursor_2(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.csem_subplot.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_csem.draw()
        else:
            self.csem_subplot.set_title("")
            self.canvas_csem.draw()


    def draw_line_2(self, event):
        if event.inaxes is not None:
            x_value = event.xdata
            if self.csem_line:
                self.csem_line.remove()
            if self.csem_annotation:
                self.csem_annotation.remove()
            self.csem_line = self.csem_subplot.axvline(x=x_value, color='g')
            y_interp = np.interp(x_value, self.csem_x_cor, self.csem_y_cor)
            self.csem_annotation = self.csem_subplot.annotate(f'({x_value:.2f}, {y_interp:.2f})', xy=(x_value, y_interp),
                                                            xytext=(x_value, y_interp))
            # Redraw the canvas
            self.canvas_csem.draw()


    # Test if the string contains only digits
    def is_num(self, ss):
        try:
            float(ss)
            return True
        except:
            return False


    def sortby(self, col, descending, current_tree):
        # Save the data of the clicked column
        data = []
        is_num = True
        for child in current_tree.get_children(''):
            data.append(
                (
                    current_tree.set(child, col),
                    child
                )
            )
            if is_num:
                if not self.is_num(data[-1][0]):
                    is_num = False
        if is_num:
            # Special handling if it's purely numeric
            data = sorted(data, key=lambda x: (float(x[0]), x[1]), reverse=descending)
        else:
            # Ignore case for sorting
            data = sorted(data, key=lambda x: (x[0].lower(), x[1]), reverse=descending)
        # Update the sorted data
        for indx, item in enumerate(data):
            current_tree.move(item[1], '', indx)
        current_tree.heading(col, command=lambda _col=col: self.sortby(col, int(not descending), current_tree))
        # Save sorting column and direction for future use,
        # allowing the original sorting to be displayed when inserting new data
        self.tree_dict[current_tree] = [col, descending]


    def update_sort(self, current_tree):
        if self.tree_dict.get(current_tree, True):
            return
        sort_column = self.tree_dict[current_tree][0]
        sort_direction = self.tree_dict[current_tree][1]
        # Apply the current sorting if it has already been set
        self.sortby(sort_column, sort_direction, current_tree)


    def delete_pro_item(self):
        # The selection() function returns a tuple containing all selected item IDs
        # These IDs are assigned to each row in your Treeview when you create it,
        # and if you haven't explicitly specified an ID during creation,
        # Tkinter will assign one automatically.
        selected_item = self.pro_tree.selection()
        if len(selected_item) == 1:
            value_1 = self.pro_tree.item(selected_item[0], "values")
            author_name = value_1[0]
            if author_name == self.master.username.get().strip():
                choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete this record?")
                if choice:
                    item_values = self.pro_tree.item(selected_item[0], "values")
                    timestamp_selected = item_values[3]
                    # Delete from json
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    for i, prox in enumerate(data["prox"]):
                        if prox["timestamp"] == timestamp_selected:
                            del data["prox"][i]
                            break
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                        json.dump(data, file, indent=4)
                    # Delete from treeview
                    self.pro_tree.delete(selected_item[0])
                    self.update_sort(self.pro_tree)
                    # Delete from word
                    doc_path_selected = f'{self.master.experiment_path_dynamic}/Metrics_Result/Proximity.docx'
                    doc = Document(doc_path_selected)
                    # Iterate through all paragraphs in the document, find and delete the corresponding Prox
                    for para in doc.paragraphs:
                        if timestamp_selected in para.text:
                            # Find the corresponding Prox and delete it
                            p = para._element
                            p.getparent().remove(p)
                    doc.save(doc_path_selected)
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    if data["prox"] == []:
                        os.remove(doc_path_selected)
            else:
                messagebox.showwarning("Warning", "Sorry, you can only delete record that you created, please select again.")


        elif len(selected_item) > 1:
            author_name_list = []
            for item_id_2 in selected_item:
                value_2 = self.pro_tree.item(item_id_2, "values")
                author_name = value_2[0]
                author_name_list.append(author_name)
            if all(name == self.master.username.get().strip() for name in author_name_list):
                choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete these records?")
                if choice:
                    # Delete from treeview, json and word
                    for item_id in selected_item:
                        item_values = self.pro_tree.item(item_id, "values")
                        timestamp_selected = item_values[3]
                        # Delete from json
                        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                            data = json.load(file)
                        for i, prox in enumerate(data["prox"]):
                            if prox["timestamp"] == timestamp_selected:
                                del data["prox"][i]
                                break
                        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                            json.dump(data, file, indent=4)
                        # Delete from treeview
                        self.pro_tree.delete(item_id)
                        self.update_sort(self.pro_tree)
                        # Delete from word
                        doc_path_selected_1 = f'{self.master.experiment_path_dynamic}/Metrics_Result/Proximity.docx'
                        doc_1 = Document(doc_path_selected_1)
                        for para in doc_1.paragraphs:
                            if timestamp_selected in para.text:
                                p = para._element
                                p.getparent().remove(p)
                        doc_1.save(doc_path_selected_1)
                        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                            data = json.load(file)
                        if data["prox"] == []:
                            os.remove(doc_path_selected_1)
            else:
                messagebox.showwarning("Warning", "Sorry, you can only delete records that you created, please select again.")


    def submit_pro_threshold(self):
        def is_float(entry):
            try:
                num = float(entry)
            except ValueError:
                return False
            return True

        if is_float(self.master.proximity_threshold.get().strip()):
            threshold = float(self.master.proximity_threshold.get().strip())
            timestamp = datetime.datetime.now().isoformat()
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            distance_list = data["distancelist"]
            total_frame = data["totalframe"]
            prox_author = self.master.username.get().strip()
            kernel = 0
            for dis in distance_list:
                if dis < threshold:
                    kernel = kernel + 1
            prox_value = kernel / total_frame
            # Insert prox result in metadata.json
            self.pro_tree.insert('', 'end', values=(prox_author, threshold, prox_value, timestamp))
            self.update_sort(self.pro_tree)
            new_prox = {
                "author": prox_author,
                "δ": threshold,
                "prox": prox_value,
                "timestamp": timestamp
            }
            data['prox'].append(new_prox)
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)
            # Create a Proximity Word document and edit it to write the new prox result
            # Add prox to metrics file as .docx
            frame_prox_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Proximity.docx'
            if os.path.exists(frame_prox_docx_path):
                doc = Document(frame_prox_docx_path)
            else:
                doc = Document()
            doc.add_paragraph(
                f'{timestamp}\n{prox_author}: \nδ = {threshold}\nProximity = {prox_value}\n--------------------------------------------------------------------------------------------------------------------')  # 在文档中添加一个段落
            doc.save(frame_prox_docx_path)

        else:
            messagebox.showwarning("Warning", "Please enter valid threshold value")


    def delete_csem_item(self):
        selected_item = self.csem_tree.selection()
        if len(selected_item) == 1:
            choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete this record?")
            if choice:
                item_values = self.csem_tree.item(selected_item[0], "values")
                timestamp_selected = item_values[4]
                # Delete from json
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                for i, csem in enumerate(data["csem"]):
                    if csem["timestamp"] == timestamp_selected:
                        del data["csem"][i]
                        break
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)
                # Delete from treeview
                self.csem_tree.delete(selected_item[0])
                self.update_sort(self.csem_tree)
                # Delete from word
                doc_path_selected = f'{self.master.experiment_path_dynamic}/Metrics_Result/CSEM.docx'
                doc = Document(doc_path_selected)
                for para in doc.paragraphs:
                    if timestamp_selected in para.text:
                        p = para._element
                        p.getparent().remove(p)
                doc.save(doc_path_selected)
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                if data["csem"] == []:
                    os.remove(doc_path_selected)


        elif len(selected_item) > 1:
            choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete these records?")
            if choice:
                # Delete from treeview, json and word
                for item_id in selected_item:
                    item_values = self.csem_tree.item(item_id, "values")
                    timestamp_selected = item_values[4]
                    # Delete from json
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    for i, csem in enumerate(data["csem"]):
                        if csem["timestamp"] == timestamp_selected:
                            del data["csem"][i]
                            break
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                        json.dump(data, file, indent=4)
                    # Delete from treeview
                    self.csem_tree.delete(item_id)
                    self.update_sort(self.csem_tree)
                    # Delete from word
                    doc_path_selected_1 = f'{self.master.experiment_path_dynamic}/Metrics_Result/CSEM.docx'
                    doc_1 = Document(doc_path_selected_1)
                    for para in doc_1.paragraphs:
                        if timestamp_selected in para.text:
                            p = para._element
                            p.getparent().remove(p)
                    doc_1.save(doc_path_selected_1)
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    if data["csem"] == []:
                        os.remove(doc_path_selected_1)


    def submit_csem_threshold(self):
        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
            data = json.load(file)
        total_frame = data["totalframe"]
        distance_list = data["distancelist"]


        def is_float(entry):
            try:
                num1 = float(entry)
            except ValueError:
                return False
            return True


        def group_max(lst, group_size):
            max_vals = []
            for i in range(0, len(lst) - group_size + 1):
                max_vals.append(max(lst[i:i + group_size]))
            return max_vals

        if is_float(self.master.csem_threshold.get().strip()):
            csem_author = self.master.username.get().strip()
            timestamp = datetime.datetime.now().isoformat()
            threshold = float(self.master.csem_threshold.get().strip())
            max_m_list = []
            for m in range(total_frame - 1):
                m = m + 1
                Nm = 0
                max_list = group_max(distance_list, m)
                for max1 in max_list:
                    if max1 < threshold:
                        Nm = Nm + 1
                if Nm > 0:
                    max_m_list.append(m)
                elif Nm == 0:
                    break
            if max_m_list != []:
                max_m_value = max_m_list[-1]
                csem_value = max_m_value / (total_frame - 1)
                self.csem_tree.insert('', 'end', values=(csem_author, threshold, max_m_value, csem_value, timestamp))
                self.update_sort(self.csem_tree)
                new_csem = {
                    "author": csem_author,
                    "δ": threshold,
                    "max(m)": max_m_value,
                    "CSEM": csem_value,
                    "timestamp": timestamp
                }
                data['csem'].append(new_csem)
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)

                frame_csem_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/CSEM.docx'
                if os.path.exists(frame_csem_docx_path):
                    doc = Document(frame_csem_docx_path)
                else:
                    doc = Document()
                doc.add_paragraph(
                    f'{timestamp}\n{csem_author}: \nδ = {threshold}\nmax(m) = {max_m_value}\nCSEM = {csem_value}\n--------------------------------------------------------------------------------------------------------------------')  # 在文档中添加一个段落
                doc.save(frame_csem_docx_path)
            else:
                self.csem_tree.insert('', 'end', values=(csem_author, threshold, 0, 0, timestamp))
                self.update_sort(self.csem_tree)
                new_csem = {
                    "author": csem_author,
                    "δ": threshold,
                    "max(m)": 0,
                    "CSEM": 0,
                    "timestamp": timestamp
                }
                data['csem'].append(new_csem)
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)

                frame_csem_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/CSEM.docx'
                if os.path.exists(frame_csem_docx_path):
                    doc = Document(frame_csem_docx_path)
                else:
                    doc = Document()
                doc.add_paragraph(
                    f'{timestamp}\n{csem_author}: \nδ = {threshold}\nmax(m) = {0}\nCSEM = {0}\n--------------------------------------------------------------------------------------------------------------------')  # 在文档中添加一个段落
                doc.save(frame_csem_docx_path)

        else:
            messagebox.showwarning("Warning", "Please enter a valid threshold.")


    def rv_di_cs(self):
        if any(os.scandir(self.master.processed_video_path_dynamic)):
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            self.master.total_frame = data['totalframe']
            tracking_method = data['trackingmethod']
            object_1_coordinate = data['object1coordinatelist']
            object_2_coordinate = data['object2coordinatelist']
            distance_list = data["distancelist"]
            total_frame_count = len(distance_list)

            self.screen_width = get_monitors()[0].width
            self.screen_height = get_monitors()[0].height

            self.rv_di_cs_toplevel = tk.Toplevel(self.master)
            self.rv_di_cs_toplevel.geometry(f"{int(self.screen_width)}x{int(self.screen_height - 200)}")

            # Set the width for each frame
            frame_width = int((self.screen_width) * 3 / 8)
            frame_height = int(self.screen_height - 200)

            # Create left frame
            self.frame_left = tk.Frame(self.rv_di_cs_toplevel, width=frame_width - 20)
            self.frame_left.pack(side='left', fill='y')

            self.rv_canvas = tk.Canvas(self.frame_left, width=int((self.screen_width) * 3 / 8) - 25)
            self.rv_scrollbar = ttk.Scrollbar(self.frame_left, orient=tk.VERTICAL,
                                                     command=self.rv_canvas.yview)
            self.rv_frame = tk.Frame(self.rv_canvas)
            self.rv_frame.bind("<Configure>", lambda e: self.rv_canvas.configure(
                scrollregion=self.rv_canvas.bbox("all")))
            self.rv_canvas.create_window((int(((self.screen_width * 3 / 8) - 25) / 2), 0),
                                                window=self.rv_frame, anchor="center")
            self.rv_canvas.configure(yscrollcommand=self.rv_scrollbar.set)
            self.rv_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.rv_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            self.rv_label_image_frame = tk.Frame(self.rv_frame, width=int(frame_width * 12 / 13),
                                                  height=int(frame_height * 1 / 4))
            self.rv_label_image_frame.pack()

            self.correlation_coefficient_label = tk.Label(self.rv_label_image_frame, text="Correlation Coefficient")
            self.correlation_coefficient_label.place(relx=0.5, rely=0.15, anchor="center")

            photo_rv = Global.photo_place("amrv.png", int((frame_width - 20) * 7 / 10), int((frame_width - 20) * 7 / 20))
            self.image_label_rv = tk.Label(self.rv_label_image_frame, image=photo_rv)
            self.image_label_rv.image = photo_rv
            self.image_label_rv.place(relx=0.5, rely=0.35, anchor="n")

            self.rv_latitude_graph_frame = tk.Frame(self.rv_frame, width=int(frame_width * 12 / 13),
                                                      height=int(frame_height * 1 / 2))
            self.rv_latitude_graph_frame.pack()

            self.rv_latitude_graph = Figure(
                figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.rv_subplot = self.rv_latitude_graph.add_subplot(111)

            object_1_latitude_list = []
            for i, coordinate_1 in enumerate(object_1_coordinate):
                object_1_latitude_list.append([i + 1, coordinate_1[1]])

            object_2_latitude_list = []
            for j, coordinate_2 in enumerate(object_2_coordinate):
                object_2_latitude_list.append([j + 1, coordinate_2[1]])

            self.object_1_x_cor = []
            self.object_1_y_cor = []
            for cor in object_1_latitude_list:
                self.object_1_x_cor.append(cor[0])
                self.object_1_y_cor.append(cor[1])

            self.object_2_x_cor = []
            self.object_2_y_cor = []
            for cor in object_2_latitude_list:
                self.object_2_x_cor.append(cor[0])
                self.object_2_y_cor.append(cor[1])

            self.rv_subplot.plot(self.object_1_x_cor, self.object_1_y_cor, color='r', label="object 1")
            self.rv_subplot.plot(self.object_2_x_cor, self.object_2_y_cor, color='b', label="object 2")
            self.rv_subplot.set_xlabel("Frame")
            self.rv_subplot.set_ylabel("Latitude (pixel)")
            # Show the legend
            self.rv_subplot.legend(frameon=True)

            # Automatically adjust the size of the chart to fit its content
            self.rv_latitude_graph.tight_layout()
            # If there is too little space between the title and
            # the top of the image, adjust the position of the subplot.
            # Adjust the top value accordingly
            self.rv_latitude_graph.subplots_adjust(top=0.9)

            # Plot the chart onto the Tkinter interface
            self.canvas_rv = FigureCanvasTkAgg(self.rv_latitude_graph,
                                                 master=self.rv_latitude_graph_frame)
            self.canvas_rv.draw()
            self.canvas_rv.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)

            self.rv_line = None
            self.rv_annotation_1 = None
            self.rv_annotation_2 = None
            self.canvas_rv.mpl_connect('button_press_event', self.draw_line_rv)
            self.canvas_rv.mpl_connect('motion_notify_event', self.move_cursor_rv)

            self.rv_scatter_graph_frame = tk.Frame(self.rv_frame, width=int(frame_width * 12 / 13),
                                                    height=int(frame_height * 1 / 2))
            self.rv_scatter_graph_frame.pack()

            self.rv_scatter_graph = Figure(
                figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.rv_subplot_2 = self.rv_scatter_graph.add_subplot(111)

            self.rv_subplot_2.scatter(self.object_1_y_cor, self.object_2_y_cor, s=5, color='r')

            self.rv_subplot_2.set_xlabel("Latitude 1 (pixel)")
            self.rv_subplot_2.set_ylabel("Latitude 2 (pixel)")

            self.rv_scatter_graph.tight_layout()
            self.rv_scatter_graph.subplots_adjust(top=0.9)

            self.canvas_rv_2 = FigureCanvasTkAgg(self.rv_scatter_graph,
                                              master=self.rv_scatter_graph_frame)
            self.canvas_rv_2.draw()
            self.canvas_rv_2.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)

            self.canvas_rv_2.mpl_connect('motion_notify_event', self.move_cursor_rv_2)

            self.rv_T_V_label_frame = tk.Frame(self.rv_frame, width=int(frame_width * 12 / 13),
                                                    height=int(frame_height * 1 / 8))
            self.rv_T_V_label_frame.pack()

            # rv T value
            self.rv_T_label = tk.Label(self.rv_T_V_label_frame, text=f"T = {self.master.total_frame}")
            self.rv_T_label.place(relx=0.5, rely=0.3, anchor='center')

            # rv V variable
            self.rv_V_label = tk.Label(self.rv_T_V_label_frame, text="V can be Latitude and Speed on this platform.")
            self.rv_V_label.place(relx=0.5, rely=0.7, anchor='center')

            self.rv_result_label_frame = tk.Frame(self.rv_frame, width=int(frame_width * 12 / 13),
                                               height=int(frame_height * 1 / 4))
            self.rv_result_label_frame.pack()

            object_1_lat_list = []
            for coordinate_5 in object_1_coordinate:
                object_1_lat_list.append(coordinate_5[1])

            object_2_lat_list = []
            for coordinate_6 in object_2_coordinate:
                object_2_lat_list.append(coordinate_6[1])

            # Calculate the correlation coefficient matrix
            pearson_coefficient_2_matrix = np.corrcoef(object_1_lat_list, object_2_lat_list)
            # Extract the correlation coefficient
            pearson_coefficient_2 = pearson_coefficient_2_matrix[0, 1]

            self.rv_Latitude_label = tk.Label(self.rv_result_label_frame, text=f"r_Latitude = {pearson_coefficient_2}")
            self.rv_Latitude_label.place(relx=0.5, rely=0.33, anchor='center')

            speed_list_1 = []
            for i in range(1, len(object_1_lat_list)):
                # Velocity is the change in position divided by the change in time (which is always 1 second here)
                speed_1 = abs(object_1_lat_list[i] - object_1_lat_list[i - 1])
                speed_list_1.append(speed_1)
            avg_speed_1 = np.mean(speed_list_1)

            speed_list_2 = []
            for j in range(1, len(object_2_lat_list)):
                speed_2 = abs(object_2_lat_list[j] - object_2_lat_list[j - 1])
                speed_list_2.append(speed_2)
            avg_speed_2 = np.mean(speed_list_2)

            # The numerator part
            numerator_4 = sum(
                (speed_list_1 - avg_speed_1) * (speed_list_2 - avg_speed_2))
            # The denominator part
            denominator_4 = np.sqrt(sum((speed_list_1 - avg_speed_1) ** 2) * sum(
                (speed_list_2 - avg_speed_2) ** 2))
            # Calculate the Pearson correlation coefficient
            pearson_coefficient_4 = numerator_4 / denominator_4

            self.rv_speed_label = tk.Label(self.rv_result_label_frame, text=f"r_Speed = {pearson_coefficient_4}")
            self.rv_speed_label.place(relx=0.5, rely=0.66, anchor='center')

            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
                data['rv'] = [{
                    "r_Latitude": pearson_coefficient_2,
                    "r_Speed": pearson_coefficient_4
                }]
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)

            frame_rv_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Correlation Coefficient.docx'
            if os.path.exists(frame_rv_docx_path):
                doc = Document(frame_rv_docx_path)
            else:
                doc = Document()
            for element in doc.element.body:
                doc.element.body.remove(element)
            doc.add_paragraph(
                f'r_Latitude = {pearson_coefficient_2}\nr_Speed = {pearson_coefficient_4}')
            doc.save(frame_rv_docx_path)

            # Create the middle frame
            self.frame_middle = tk.Frame(self.rv_di_cs_toplevel, width=frame_width - 20)
            self.frame_middle.pack(side='left', fill='y')

            self.di_canvas = tk.Canvas(self.frame_middle, width=int((self.screen_width) * 3 / 8) - 25)
            self.di_scrollbar = ttk.Scrollbar(self.frame_middle, orient=tk.VERTICAL,
                                              command=self.di_canvas.yview)
            self.di_frame = tk.Frame(self.di_canvas)
            self.di_frame.bind("<Configure>", lambda e: self.di_canvas.configure(
                scrollregion=self.di_canvas.bbox("all")))
            self.di_canvas.create_window((int(((self.screen_width * 3 / 8) - 25) / 2), 0),
                                         window=self.di_frame, anchor="center")
            self.di_canvas.configure(yscrollcommand=self.di_scrollbar.set)
            self.di_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.di_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            self.di_label_image_frame = tk.Frame(self.di_frame, width=int(frame_width * 12 / 13),
                                                 height=int(frame_height * 2 / 5))
            self.di_label_image_frame.pack()

            self.dynamic_interaction_label = tk.Label(self.di_label_image_frame, text="Dynamic Interaction Index")
            self.dynamic_interaction_label.place(relx=0.5, rely=0.1, anchor="center")

            photo_di = Global.photo_place("amDI.png", int((frame_width - 20) * 9 / 10), int((frame_width - 20) * 9 / 20))
            self.image_label_di = tk.Label(self.di_label_image_frame, image=photo_di)
            self.image_label_di.image = photo_di
            self.image_label_di.place(relx=0.5, rely=0.2, anchor="n")

            self.di_displacement_graph_frame = tk.Frame(self.di_frame, width=int(frame_width * 12 / 13),
                                                    height=int(frame_height * 1 / 2))
            self.di_displacement_graph_frame.pack()

            self.di_displacement_graph = Figure(
                figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.di_subplot = self.di_displacement_graph.add_subplot(111)

            object_1_position_list = []
            for coordinate_7 in object_1_coordinate:
                object_1_position_list.append(coordinate_7[1])

            object_2_position_list = []
            for coordinate_8 in object_2_coordinate:
                object_2_position_list.append(coordinate_8[1])

            self.object_1_displacement_list = [abs(object_1_position_list[i + 1] - object_1_position_list[i]) for i in range(len(object_1_position_list) - 1)]
            self.object_2_displacement_list = [abs(object_2_position_list[i + 1] - object_2_position_list[i]) for i in range(len(object_2_position_list) - 1)]

            self.object_1_index_list = [i + 1 for i in range(len(self.object_1_displacement_list))]
            self.object_2_index_list = [i + 1 for i in range(len(self.object_2_displacement_list))]

            self.di_subplot.plot(self.object_1_index_list, self.object_1_displacement_list, color='r', label="object 1")
            self.di_subplot.plot(self.object_2_index_list, self.object_2_displacement_list, color='b', label="object 2")
            self.di_subplot.set_xlabel("Frame")
            self.di_subplot.set_ylabel("Displacement Between Frames (pixel)")
            self.di_subplot.legend(frameon=True)

            self.di_displacement_graph.tight_layout()
            self.di_displacement_graph.subplots_adjust(top=0.9)

            self.canvas_di = FigureCanvasTkAgg(self.di_displacement_graph,
                                               master=self.di_displacement_graph_frame)
            self.canvas_di.draw()
            self.canvas_di.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)

            self.di_line = None
            self.di_annotation_1 = None
            self.di_annotation_2 = None
            self.canvas_di.mpl_connect('button_press_event', self.draw_line_di)
            self.canvas_di.mpl_connect('motion_notify_event', self.move_cursor_di)

            self.di_angle_graph_frame = tk.Frame(self.di_frame, width=int(frame_width * 12 / 13),
                                                        height=int(frame_height * 1 / 2))
            self.di_angle_graph_frame.pack()

            self.di_angle_graph = Figure(
                figsize=(int(frame_width * 12 / 13) / 100, int(frame_height * 1 / 2) / 100), dpi=100)
            self.di_subplot_2 = self.di_angle_graph.add_subplot(111)

            # Example sequence
            object_1_displacement_neg_list = [object_1_position_list[i + 1] - object_1_position_list[i] for i in
                                              range(len(object_1_position_list) - 1)]
            object_2_displacement_neg_list = [object_2_position_list[i + 1] - object_2_position_list[i] for i in
                                              range(len(object_2_position_list) - 1)]

            object_1_angle_list = [90 if y < 0 else (-90 if y > 0 else 0) for y in object_1_displacement_neg_list]
            object_2_angle_list = [90 if y < 0 else (-90 if y > 0 else 0) for y in object_2_displacement_neg_list]

            object_1_angle_frame_list = [i + 1 for i in range(len(object_1_angle_list))]
            object_2_angle_frame_list = [i + 1 for i in range(len(object_2_angle_list))]

            self.di_subplot_2.scatter(object_1_angle_frame_list, object_1_angle_list, s=5, color='r', label="object 1")
            self.di_subplot_2.scatter(object_2_angle_frame_list, object_2_angle_list, s=5, color='b', label="object 2")
            self.di_subplot_2.set_xlabel("Frame")
            self.di_subplot_2.set_ylabel("Angle (degree)")
            self.di_subplot_2.legend(frameon=True)

            self.di_angle_graph.tight_layout()
            self.di_angle_graph.subplots_adjust(top=0.9)

            self.canvas_di_2 = FigureCanvasTkAgg(self.di_angle_graph,
                                               master=self.di_angle_graph_frame)
            self.canvas_di_2.draw()
            self.canvas_di_2.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)

            self.canvas_di_2.mpl_connect('motion_notify_event', self.move_cursor_di_2)

            self.di_T_dis_theta_label_frame = tk.Frame(self.di_frame, width=int(frame_width * 12 / 13),
                                               height=int(frame_height * 1 / 4))
            self.di_T_dis_theta_label_frame.pack()

            # di T value
            self.di_T_label = tk.Label(self.di_T_dis_theta_label_frame, text=f"T = {self.master.total_frame}")
            self.di_T_label.place(relx=0.5, rely=0.2, anchor='center')

            # di angle variable
            self.di_angle_label = tk.Label(self.di_T_dis_theta_label_frame,
                                           text="θ can only be equal to 0, 90, -90 degrees on this platform.")
            self.di_angle_label.place(relx=0.5, rely=0.8, anchor='center')

            self.di_result_label_frame = tk.Frame(self.di_frame, width=int(frame_width * 12 / 13),
                                               height=int(frame_height * 1 / 4))
            self.di_result_label_frame.pack()

            T_variable = len(self.object_1_displacement_list)
            total_sum_di_dis = 0
            for t in range(T_variable):
                denominator = self.object_1_displacement_list[t] + self.object_2_displacement_list[t]
                if denominator != 0:
                    total_sum_di_dis += (1 - abs(
                        self.object_1_displacement_list[t] - self.object_2_displacement_list[t]) / denominator)
                else:
                    total_sum_di_dis += 1
            # Avoid division by 0
            DI_d = total_sum_di_dis / T_variable if T_variable != 0 else 0

            self.di_displacement_label = tk.Label(self.di_result_label_frame, text=f"DI_d = {DI_d}")
            self.di_displacement_label.place(relx=0.5, rely=0.25, anchor='center')

            theta_diff_list = [a - b for a, b in zip(object_1_angle_list, object_2_angle_list)]

            total_sum_di_theta = 0
            for t in range(T_variable):
                absolute_angle = round(math.cos(math.radians(theta_diff_list[t])), 2)
                total_sum_di_theta += absolute_angle
            DI_theta = total_sum_di_theta / T_variable

            self.di_theta_label = tk.Label(self.di_result_label_frame, text=f"DI_θ = {DI_theta}")
            self.di_theta_label.place(relx=0.5, rely=0.5, anchor='center')

            total_sum_di = 0
            for t in range(T_variable):
                # Sum of displacement differences
                displacement_sum = self.object_1_displacement_list[t] + self.object_2_displacement_list[t]
                term1 = round(math.cos(math.radians(theta_diff_list[t])), 2)
                # Exclude cases where the divisor is 0
                if displacement_sum != 0:
                    term2 = 1 - (abs(
                        self.object_1_displacement_list[t] - self.object_2_displacement_list[t])) / displacement_sum
                else:
                    term2 = 1
                total_sum_di += term1 * term2
            # Avoid division by 0
            DI_com = total_sum_di / T_variable if T_variable != 0 else 0

            self.di_com_label = tk.Label(self.di_result_label_frame, text=f"DI = {DI_com}")
            self.di_com_label.place(relx=0.5, rely=0.75, anchor='center')

            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
                data['di'] = [{
                    "DI_d": DI_d,
                    "DI_θ": DI_theta,
                    "DI": DI_com
                }]
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)

            frame_di_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Dynamic Interaction.docx'
            if os.path.exists(frame_di_docx_path):
                doc = Document(frame_di_docx_path)
            else:
                doc = Document()
            for element in doc.element.body:
                doc.element.body.remove(element)
            doc.add_paragraph(
                f'DI_d = {DI_d}\nDI_θ = {DI_theta}\nDI = {DI_com}')
            doc.save(frame_di_docx_path)

            # Create right frame
            self.frame_right = tk.Frame(self.rv_di_cs_toplevel, width=int(frame_width * 2 / 3))
            self.frame_right.pack(side='left', fill='both', expand=True)

            self.cs_label_2 = tk.Label(self.frame_right, text="Coefficient of Sociality (Cs)")
            self.cs_label_2.place(relx=0.5, rely=0.05, anchor='center')

            photo_right = Global.photo_place("amCs.png", int(frame_width / 2), int(frame_width / 3))
            self.image_label_right = tk.Label(self.frame_right, image=photo_right)
            self.image_label_right.image = photo_right
            self.image_label_right.place(relx=0.5, rely=0.1, anchor='n')

            # cs T value
            self.cs_T_label_coor = tk.Label(self.frame_right, text=f"T = {total_frame_count}")
            self.cs_T_label_coor.place(relx=0.5, rely=0.4, anchor='center')

            if tracking_method == "Multiple object tracking":
                width_ratio = Global.multiple_width_ratio
                height_ratio = Global.multiple_height_ratio
            elif tracking_method == "Blob tracking":
                width_ratio = Global.blob_width_ratio
                height_ratio = Global.blob_height_ratio

            all_distance_list = []
            for cor1 in object_1_coordinate:
                for cor2 in object_2_coordinate:
                    distance = round(math.sqrt(((cor1[0] - cor2[0]) / width_ratio) ** 2 + ((cor1[1] - cor2[1]) / height_ratio) ** 2), 2)
                    all_distance_list.append(distance)

            # cs DO value
            DO_value = sum(distance_list) / total_frame_count
            self.cs_do_label_coor = tk.Label(self.frame_right, text=f"D_O = {DO_value}")
            self.cs_do_label_coor.place(relx=0.5, rely=0.45, anchor='center')

            # cs DE value
            DE_value = sum(all_distance_list) / total_frame_count ** 2
            self.cs_de_label_coor = tk.Label(self.frame_right, text=f"D_E = {DE_value}")
            self.cs_de_label_coor.place(relx=0.5, rely=0.5, anchor='center')

            # cs Cs value
            Cs_value = (DE_value - DO_value) / (DE_value + DO_value)
            self.cs_cs_label_coor = tk.Label(self.frame_right, text=f"Cs = {Cs_value}")
            self.cs_cs_label_coor.place(relx=0.5, rely=0.55, anchor='center')

            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
                data['cs'] = [{
                    "D_O": DO_value,
                    "D_E": DE_value,
                    "Cs": Cs_value
                }]
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)

            frame_cs_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Cs.docx'
            if os.path.exists(frame_cs_docx_path):
                doc = Document(frame_cs_docx_path)
            else:
                doc = Document()
            for element in doc.element.body:
                doc.element.body.remove(element)
            doc.add_paragraph(
                f'D_O = {DO_value}\nD_E = {DE_value}\nCs = {Cs_value}')
            doc.save(frame_cs_docx_path)

            def move_scrollbar_to_top():
                self.rv_canvas.yview_moveto(0)
                self.di_canvas.yview_moveto(0)
            self.master.after(100, move_scrollbar_to_top)

        else:
            messagebox.showwarning("Warning", "Please process the video first")


    def draw_line_di(self, event):
        if event.inaxes is not None:
            x_value = event.xdata
            if self.di_line in self.di_subplot.lines:
                self.di_line.remove()
            if self.di_annotation_1 in self.di_subplot.texts:
                self.di_annotation_1.remove()
            if self.di_annotation_2 in self.di_subplot.texts:
                self.di_annotation_2.remove()

            self.di_line = self.di_subplot.axvline(x=x_value, color='g')
            y_value_1 = np.interp(x_value, self.object_1_index_list, self.object_1_displacement_list)
            y_value_2 = np.interp(x_value, self.object_2_index_list, self.object_2_displacement_list)
            self.di_annotation_1 = self.di_subplot.annotate(f'({x_value:.2f}, {y_value_1:.2f})', xy=(x_value, y_value_1),
                                                          xytext=(x_value, y_value_1))
            self.di_annotation_2 = self.di_subplot.annotate(f'({x_value:.2f}, {y_value_2:.2f})', xy=(x_value, y_value_2),
                                                          xytext=(x_value, y_value_2))
            self.canvas_di.draw()


    def move_cursor_di(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.di_subplot.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_di.draw()
        else:
            self.di_subplot.set_title("")
            self.canvas_di.draw()


    def move_cursor_di_2(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.di_subplot_2.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_di_2.draw()
        else:
            self.di_subplot_2.set_title("")
            self.canvas_di_2.draw()


    def draw_line_rv(self, event):
        if event.inaxes is not None:
            x_value = event.xdata
            if self.rv_line in self.rv_subplot.lines:
                self.rv_line.remove()
            if self.rv_annotation_1 in self.rv_subplot.texts:
                self.rv_annotation_1.remove()
            if self.rv_annotation_2 in self.rv_subplot.texts:
                self.rv_annotation_2.remove()

            self.rv_line = self.rv_subplot.axvline(x=x_value, color='g')
            y_value_1 = np.interp(x_value, self.object_1_x_cor, self.object_1_y_cor)
            y_value_2 = np.interp(x_value, self.object_2_x_cor, self.object_2_y_cor)
            self.rv_annotation_1 = self.rv_subplot.annotate(f'({x_value:.2f}, {y_value_1:.2f})', xy=(x_value, y_value_1),
                                                          xytext=(x_value, y_value_1))
            self.rv_annotation_2 = self.rv_subplot.annotate(f'({x_value:.2f}, {y_value_2:.2f})', xy=(x_value, y_value_2),
                                                          xytext=(x_value, y_value_2))
            self.canvas_rv.draw()


    def move_cursor_rv(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.rv_subplot.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_rv.draw()
        else:
            self.rv_subplot.set_title("")
            self.canvas_rv.draw()


    def move_cursor_rv_2(self, event):
        if event.inaxes is not None:
            x = event.xdata
            y = event.ydata
            self.rv_subplot_2.set_title(f'x={x:.2f}, y={y:.2f}')
            self.canvas_rv_2.draw()
        else:
            self.rv_subplot_2.set_title("")
            self.canvas_rv_2.draw()