import json
from timeit import Timer
from tkinter import messagebox
import cv2 as cv
import numpy as np


def code_to_test():
    import threading
    import time
    import tkinter as tk
    from tkinter import ttk
    from tkinter import messagebox
    import os
    import json
    import datetime
    import math
    import PIL.Image
    import PIL.ImageTk
    from PIL import ImageEnhance
    from screeninfo import get_monitors
    from docx import Document
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    from matplotlib.figure import Figure
    import glob
    import Global

    # Processing Page frame class
    class ProcessingPage(tk.Frame):
        def __init__(self, master):
            tk.Frame.__init__(self, master)

            # 保存tree的排序数据指标
            self.tree_dict = dict()
            self.start_processing_button = tk.Button(self, text="Start Processing", command=self.start_processing,
                                                     cursor="hand2")
            self.start_processing_button.place(relx=0.5, rely=0.25, anchor='center')
            self.start_processing_label = tk.Label(self, text=self.master.process_status)
            self.start_processing_label.place(relx=0.5, rely=0.4, anchor='center')
            self.analyze_metrics_btn = tk.Button(self, text="Analyze with Metrics", command=self.analyze_with_metrics,
                                                 cursor="hand2")
            self.analyze_metrics_btn.place(relx=0.5, rely=0.55, anchor='center')
            self.review_button = tk.Button(self, text="Review and Make Notes", command=self.review_toplevel,
                                           cursor="hand2")
            self.review_button.place(relx=0.5, rely=0.7, anchor='center')
            self.result_button = tk.Button(self, text="View the Results", command=self.open_result_toplevel,
                                           cursor="hand2")
            self.result_button.place(relx=0.5, rely=0.85, anchor='center')
            self.back_button = tk.Button(self, text="Back", command=self.back_to_upload_page_or_experiment_page,
                                         cursor="hand2")
            self.back_button.place(relx=0.05, rely=0.05, anchor='nw')

            self.master.protocol("WM_DELETE_WINDOW", self.quit)

        # 用抛出来测试字符串里是不是纯数字
        def is_num(self, ss):
            try:
                float(ss)
                return True
            except:
                return False

        def sortby(self, col, descending, current_tree):
            # 保存点击的列的数据
            data = []
            is_num = True
            for child in current_tree.get_children(''):
                data.append(
                    (
                        current_tree.set(child, col),
                        child
                    )
                )
                if is_num:
                    if not self.is_num(data[-1][0]):
                        is_num = False
            if is_num:
                # 如果是纯数字特殊处理
                data = sorted(data, key=lambda x: (float(x[0]), x[1]), reverse=descending)
            else:
                # 无视大小写排序
                data = sorted(data, key=lambda x: (x[0].lower(), x[1]), reverse=descending)
            # 更新排序好的数据
            for indx, item in enumerate(data):
                current_tree.move(item[1], '', indx)
            current_tree.heading(col, command=lambda _col=col: self.sortby(col, int(not descending), current_tree))
            # 保存排序列和方向, 方便下次插入新数据时, 按照原先的排序进行展示
            self.tree_dict[current_tree] = [col, descending]

        def update_sort(self, current_tree):
            # 根据传递
            if self.tree_dict.get(current_tree, True):
                return
            sort_column = self.tree_dict[current_tree][0]
            sort_direction = self.tree_dict[current_tree][1]
            # 如果已经设置了排序，重新应用当前的排序
            self.sortby(sort_column, sort_direction, current_tree)

        def analyze_with_metrics(self):
            from MetricsSelectionPage import MetricsSelectionPage
            if self.master.process_status == "Processed":
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                if data['boxes'] == 2:
                    self.master.show_frame(MetricsSelectionPage)
                else:
                    messagebox.showwarning("Warning",
                                           f"Sorry, this platform can only analyze dyadic movement.\nYou have selected {data['boxes']} ROI(s). If you wish to view the analysis results,\nplease process again and select two objects as ROIs.")
            else:
                messagebox.showwarning("Warning", "Please process the video first")

        def start_processing(self):
            from TrackingSelectionPage import TrackingSelectionPage
            if self.master.process_status == "Processed":
                choice = messagebox.askyesno("Confirmation",
                                             "The source video has been processed, do you want to process again?\nIf you do that, all your annotations and metrics records will be deleted!")
                if choice:
                    # 读取 JSON 文件
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    # 对数据进行更新
                    self.master.process_status = "Unprocessed"
                    data['status'] = "Unprocessed"
                    data['totalframe'] = 0
                    data['distancelist'] = []
                    data['object1coordinatelist'] = []
                    data['object2coordinatelist'] = []
                    data['prox'] = []
                    data['csem'] = []
                    data['cs'] = []
                    data['rv'] = []
                    data['di'] = []
                    data['annotation'] = []
                    data['boxes'] = 0
                    data['trackingmethod'] = ''
                    # 将更新后的数据写回文件
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                        json.dump(data, file, indent=4)
                    # 删除processed video
                    dir = f'{self.master.experiment_path_dynamic}/Processed_Video'
                    for files in os.listdir(dir):
                        path = os.path.join(dir, files)
                        os.remove(path)
                    # 找到Annotation文件夹下的所有Word文档，无论是 .doc 还是 .docx
                    word_files_anno = glob.glob(
                        os.path.join(f'{self.master.experiment_path_dynamic}/Annotation', "*.doc*"))
                    for word_file_anno in word_files_anno:
                        # 删除指定的 Word 文档
                        os.remove(word_file_anno)
                    # 找到Metrics_Result文件夹下的所有Word文档，无论是 .doc 还是 .docx
                    word_files_metrics = glob.glob(
                        os.path.join(f'{self.master.experiment_path_dynamic}/Metrics_Result', "*.doc*"))
                    for word_file_metrics in word_files_metrics:
                        # 删除指定的 Word 文档
                        os.remove(word_file_metrics)

                    self.master.show_frame(TrackingSelectionPage)
            else:
                self.master.show_frame(TrackingSelectionPage)

        def review_toplevel(self):
            if any(os.scandir(self.master.processed_video_path_dynamic)):

                self.screen_width = get_monitors()[0].width
                self.screen_height = get_monitors()[0].height

                self.review_toplevel = tk.Toplevel(self.master)
                self.review_toplevel.geometry(f"{int(self.screen_width - 200)}x{int(self.screen_height - 150)}")

                # 创建左侧Frame，并添加一个Label
                self.left_frame = tk.Frame(self.review_toplevel)
                self.left_frame.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)

                self.cap = cv.VideoCapture(
                    f'{self.master.processed_video_path_dynamic}/{self.master.basement_name_dynamic}')

                def on_closing():
                    if self.cap.isOpened():
                        self.cap.release()
                    self.master.thread_running = False
                    self.review_toplevel.destroy()

                self.review_toplevel.protocol("WM_DELETE_WINDOW", on_closing)
                # 在左侧Frame中创建一个内嵌的Frame来装视频的Label和Scale
                self.video_scrollbar = tk.Frame(self.left_frame)
                self.video_scrollbar.place(relx=0.5, rely=0.47, relwidth=0.9, relheight=0.89, anchor='center')

                self.movieLabel = tk.Label(self.video_scrollbar)  # 创建一个用于播放视频的label容器
                self.movieLabel.pack()

                self.pause_video = True  # This variable will control when to pause video_player

                self.slider = tk.Scale(self.video_scrollbar, from_=1, to=self.cap.get(cv.CAP_PROP_FRAME_COUNT),
                                       orient=tk.HORIZONTAL, command=self.slider_moved)
                self.slider.pack(padx=20, pady=(20, 0), anchor='s', fill=tk.X)

                self.play_pause_btn = tk.Button(self.left_frame, text='Play', command=self.play_pause, cursor="hand2")
                self.play_pause_btn.place(relx=0.1, rely=0.95, anchor='w')
                self.annotation_btn = tk.Button(self.left_frame, text='Make Notes', command=self.make_notes,
                                                cursor="hand2")
                self.annotation_btn.place(relx=0.9, rely=0.95, anchor='e')

                # 创建右侧Frame
                self.right_frame = tk.Frame(self.review_toplevel)
                self.right_frame.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)

                # 在右侧Frame中创建一个内嵌的Frame来装Treeview和Scrollbar
                self.annotation_tree_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
                self.annotation_tree_frame.place(relx=0.5, rely=0.275, relwidth=0.9, relheight=0.5, anchor='center')

                self.annotation_tree = ttk.Treeview(self.annotation_tree_frame, columns=("c1", "c2", "c3", "c4"),
                                                    show='headings')
                self.annotation_tree.column("c1", anchor=tk.CENTER, width=25)
                self.annotation_tree.heading("c1", text="Frame",
                                             command=lambda _col="c1": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c2", anchor=tk.CENTER, width=50)
                self.annotation_tree.heading("c2", text="Author",
                                             command=lambda _col="c2": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c3", anchor=tk.CENTER, width=325)
                self.annotation_tree.heading("c3", text="Notes",
                                             command=lambda _col="c3": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c4", width=0, minwidth=0, stretch='no')
                self.annotation_tree.heading("c4", text="Timestamp",
                                             command=lambda _col="c4": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

                self.annotation_tree.bind('<<TreeviewSelect>>', self.on_treeview_select)
                # 创建一个字典用来保存每个 Treeview 项目的完整内容
                self.full_contents = {}

                # Create a vertical scrollbar
                self.annotation_tree_vsb = ttk.Scrollbar(self.annotation_tree_frame, orient="vertical",
                                                         command=self.annotation_tree.yview)
                self.annotation_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
                # Configure the treeview to use the scrollbar
                self.annotation_tree.configure(yscrollcommand=self.annotation_tree_vsb.set)

                # Add annotation to treeview from metadata.json
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)

                for note in data["annotation"]:
                    current_frame = note["frame"]
                    annotation_author = note["author"]
                    text_content = note["text"]
                    timestamp = note["timestamp"]
                    split_text_content = text_content.split("\n")[0] if "\n" in text_content else text_content
                    # Add annotation to treeview
                    item_id = self.annotation_tree.insert('', 'end', values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.full_contents[item_id] = text_content
                self.update_sort(self.annotation_tree)

                # 创建容纳Text的Frame
                self.text_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
                self.text_frame.place(relx=0.495, rely=0.73, relwidth=0.9, relheight=0.36, anchor='center')

                # 创建Text
                self.display_text = tk.Text(self.text_frame, width=40, height=10)
                self.display_text.pack(fill=tk.BOTH, expand=True)

                # 创建manage my own annotations button
                self.manage_own_btn = tk.Button(self.right_frame, text="Manage My Own Annotations",
                                                command=self.manage_own_annotations, cursor="hand2")
                self.manage_own_btn.place(relx=0.5, rely=0.95, anchor='center')
                self.master.thread_running = True
                threading.Thread(target=self.video_player).start()

                self.master.protocol("WM_DELETE_WINDOW", self.on_closing)

            else:
                messagebox.showwarning("Warning", "Please process the video first")

        def on_closing(self):
            self.master.thread_running = False  # 设置thread_running为False
            self.master.destroy()  # 关闭窗口
            os._exit(0)

        def open_result_toplevel(self):
            if any(os.scandir(self.master.processed_video_path_dynamic)):
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                if data['boxes'] == 2:
                    self.result_toplevel_window()
                else:
                    messagebox.showwarning("Warning",
                                           f"Sorry, this platform can only analyze dyadic movement.\nYou have selected {data['boxes']} ROI(s). If you wish to view the analysis results,\nplease process again and select two objects as ROIs.")
            else:
                messagebox.showwarning("Warning", "Please process the video first")

        def result_toplevel_window(self):
            self.screen_width = get_monitors()[0].width
            self.screen_height = get_monitors()[0].height
            self.result_toplevel = tk.Toplevel(self.master)
            self.result_toplevel.geometry(f"{int(self.screen_width - 170)}x{int(self.screen_height - 150)}")
            self.result_toplevel.resizable(False, False)
            toplevel_width = int(self.screen_width - 200)
            toplevel_height = int(self.screen_height - 150)
            half_width = toplevel_width // 2
            self.frame_container = tk.Frame(self.result_toplevel)
            self.frame_container.pack(fill="both", expand=True)
            # 创建左侧Frame，并添加一个Label
            self.left_frame_1 = tk.Frame(self.frame_container, width=half_width)
            self.left_frame_1.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
            self.cap_1 = cv.VideoCapture(
                f'{self.master.processed_video_path_dynamic}/{self.master.basement_name_dynamic}')

            def on_closing():
                if self.cap_1.isOpened():
                    self.cap_1.release()
                self.master.thread_running = False
                self.result_toplevel.destroy()

            self.result_toplevel.protocol("WM_DELETE_WINDOW", on_closing)
            # 在左侧Frame中创建一个内嵌的Frame来装视频的Label和Scale
            self.video_scrollbar_1 = tk.Frame(self.left_frame_1)
            self.video_scrollbar_1.place(relx=0.5, rely=0.47, relwidth=0.9, relheight=0.89, anchor='center')
            width_video_scrollbar_1 = self.video_scrollbar_1.winfo_width()
            height_video_scrollbar_1 = self.video_scrollbar_1.winfo_height()
            self.video_frame_1 = tk.Frame(self.video_scrollbar_1, bg="red", width=width_video_scrollbar_1,
                                          height=int(height_video_scrollbar_1 * 1 / 8))
            self.video_frame_1.pack()
            self.movieLabel_1 = tk.Label(self.video_frame_1)  # 创建一个用于播放视频的label容器
            self.movieLabel_1.pack(fill=tk.BOTH, expand=1)
            # self.movieLabel.place(relx=0.08, rely=0.05, anchor='nw')
            self.pause_video_1 = True  # This variable will control when to pause video_player
            self.slider_1 = tk.Scale(self.video_scrollbar_1, from_=1, to=self.cap_1.get(cv.CAP_PROP_FRAME_COUNT),
                                     orient=tk.HORIZONTAL, command=self.slider_moved_1)
            self.slider_1.pack(padx=20, pady=(20, 0), anchor='s', fill=tk.X)
            self.play_pause_btn_1 = tk.Button(self.left_frame_1, text='Play', command=self.play_pause_1, cursor="hand2")
            self.play_pause_btn_1.place(relx=0.1, rely=0.95, anchor='w')
            self.master.thread_running = True
            threading.Thread(target=self.video_player_1).start()
            # 创建右侧Frame
            self.right_frame_1 = tk.Frame(self.frame_container, width=half_width)
            self.right_frame_1.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)
            # Create a canvas
            self.canvas = tk.Canvas(self.right_frame_1, width=half_width - 25)
            self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)
            # Add a scrollbar to the canvas
            self.full_scrollbar = ttk.Scrollbar(self.right_frame_1, orient=tk.VERTICAL, command=self.canvas.yview)
            self.full_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the canvas
            self.canvas.configure(yscrollcommand=self.full_scrollbar.set)
            self.canvas.bind('<Configure>', lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
            # Create a container frame inside the canvas
            self.container_canvas_frame = tk.Frame(self.canvas)
            quarter_width = half_width // 2
            # Add that new frame to a window in the canvas
            self.canvas.create_window((quarter_width, 0), window=self.container_canvas_frame, anchor="center")
            # 然而，与Frame或其他容器部件不同，Canvas本身并不直接支持添加其他部件。你不能像在Frame中那样直接使用.pack()、.grid()或.place()方法来在Canvas中添加一个部件。
            # 因此，要在Canvas中添加一个部件（在这个例子中是Frame），你需要使用create_window()方法。
            # canvas.create_window((0, 0), window=inner_frame, anchor="nw")这行代码的作用就是在Canvas上创建一个窗口，
            # 该窗口的位置设置为 (0, 0)，窗口的内容是 inner_frame，并将其锚点设为 'nw'。这意味着，该窗口将会以其左上角（northwest）为参照点，
            # 位于Canvas的 (0, 0) 位置。这样，你就可以在这个窗口（实质上是在Canvas上的一块区域）中添加其他的部件，例如Label、Button等。
            self.optimus_prime = tk.Frame(self.container_canvas_frame, width=half_width - 20)
            self.optimus_prime.pack()
            self.annotation_container = tk.Frame(self.container_canvas_frame, width=int(half_width * 12 / 13),
                                                 height=int(toplevel_height * 2 / 5))
            self.annotation_container.pack()
            self.annotation_all_tree_frame = tk.Frame(
                self.annotation_container)  # Create a frame for the Treeview and Scrollbar
            self.annotation_all_tree_frame.place(relx=0.5, rely=0.3, relwidth=1, relheight=0.5, anchor='center')
            self.annotation_all_tree = ttk.Treeview(self.annotation_all_tree_frame, columns=("c1", "c2", "c3", "c4"),
                                                    show='headings', selectmode="extended", height=5)
            self.annotation_all_tree.column("c1", anchor=tk.CENTER, width=25)
            self.annotation_all_tree.heading("c1", text="Frame",
                                             command=lambda _col="c1": self.sortby(_col, 0, self.annotation_all_tree))
            self.annotation_all_tree.column("c2", anchor=tk.CENTER, width=50)
            self.annotation_all_tree.heading("c2", text="Author",
                                             command=lambda _col="c2": self.sortby(_col, 0, self.annotation_all_tree))
            self.annotation_all_tree.column("c3", anchor=tk.CENTER, width=325)
            self.annotation_all_tree.heading("c3", text="Notes",
                                             command=lambda _col="c3": self.sortby(_col, 0, self.annotation_all_tree))
            self.annotation_all_tree.column("c4", width=0, minwidth=0, stretch='no')
            self.annotation_all_tree.heading("c4", text="Timestamp",
                                             command=lambda _col="c4": self.sortby(_col, 0, self.annotation_all_tree))
            self.annotation_all_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            self.annotation_all_tree.bind('<<TreeviewSelect>>', self.on_all_treeview_select)
            # 创建一个字典用来保存每个 Treeview 项目的完整内容
            self.all_full_contents = {}
            # Create a vertical scrollbar
            self.annotation_all_tree_vsb = ttk.Scrollbar(self.annotation_all_tree_frame, orient="vertical",
                                                         command=self.annotation_all_tree.yview)
            self.annotation_all_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.annotation_all_tree.configure(yscrollcommand=self.annotation_all_tree_vsb.set)
            # 创建Text
            self.display_all_text = tk.Text(self.annotation_container)
            self.display_all_text.place(relx=0.5, rely=0.78, relwidth=1, relheight=0.35, anchor='center')
            # Add annotation to treeview from metadata.json
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)

            def insert_sorted(tree, frame, author, text, timestamp):
                all_items = tree.get_children()
                index = 0
                for index, item_id in enumerate(all_items):
                    item_frame = tree.item(item_id, 'values')[0]
                    if int(frame) < int(item_frame):
                        break
                else:
                    index += 1
                return tree.insert('', index, values=(frame, author, text, timestamp))

            for note in data["annotation"]:
                current_frame = note["frame"]
                annotation_author = note["author"]
                text_content = note["text"]
                timestamp = note["timestamp"]
                split_text_content = text_content.split("\n")[0] if "\n" in text_content else text_content
                # Add annotation to treeview sorted by frame
                all_item_id_list = self.annotation_all_tree.get_children()
                # 在treeview.insert方法中，index参数是一个指标，用来表示新项将会被插入的位置。
                # 比如，index=0表示新项将会被插入到所有项的最前面，成为第一个项。
                # index=1表示新项将会被插入到第一个项的后面，成为第二个项。
                # 如果index等于treeview中项的数量，新项将会被插入到所有项的后面，成为最后一个项。
                if not all_item_id_list:
                    item_id = self.annotation_all_tree.insert("", 0, values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.all_full_contents[item_id] = text_content
                else:
                    index = 0
                    for i_id in all_item_id_list:
                        i_value = self.annotation_all_tree.item(i_id)
                        i_frame = i_value['values'][0]
                        # If the current frame is less than the item frame, we found our index
                        if int(current_frame) < int(i_frame):
                            break
                        index = index + 1  # Increase index by 1 for each loop
                    item_id = self.annotation_all_tree.insert("", index, values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.all_full_contents[item_id] = text_content
            self.update_sort(self.annotation_all_tree)
            self.graph_container = tk.Frame(self.container_canvas_frame, width=int(half_width * (12 / 13)))
            self.graph_container.pack()
            self.graph_1_frame = tk.Frame(self.container_canvas_frame, width=int(half_width * 12 / 13),
                                          height=int(toplevel_height * 3 / 5))
            self.graph_1_frame.pack()
            self.inverse_relative_position_label = tk.Label(self.graph_1_frame,
                                                            text="Inverse Relative Position Graph (mm)")
            self.inverse_relative_position_label.place(relx=0.5, rely=0.08, anchor='center')
            self.inverse_relative_position_graph_frame = tk.Frame(self.graph_1_frame)
            self.inverse_relative_position_graph_frame.place(relx=0.5, rely=0.55, relwidth=1, relheight=0.8,
                                                             anchor='center')
            self.inverse_relative_position_graph = Figure()
            self.subplot_1 = self.inverse_relative_position_graph.add_subplot(111)
            # 读取 JSON 文件
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            tracking_method = data['trackingmethod']
            total_frame = data["totalframe"]
            self.total_frame = data["totalframe"]
            obj_1_cor = data["object1coordinatelist"]
            obj_2_cor = data["object2coordinatelist"]
            if tracking_method == "Multiple object tracking":
                width_ratio = Global.multiple_width_ratio
                height_ratio = Global.multiple_height_ratio
            elif tracking_method == "Blob tracking":
                width_ratio = Global.blob_width_ratio
                height_ratio = Global.blob_height_ratio
            origin_y_1 = obj_1_cor[0][1]
            origin_y_2 = obj_2_cor[0][1]
            self.x_cor_1 = [(frame + 1) for frame in range(total_frame)]
            # 实际距离毫米mm
            self.y_cor_1_1 = [((origin_y_1 - obj_1_cor[frame_add - 1][1]) / height_ratio * 10) for frame_add in
                              self.x_cor_1]
            self.y_cor_2_1 = [((origin_y_2 - obj_2_cor[frame_add - 1][1]) / height_ratio * 10) for frame_add in
                              self.x_cor_1]
            self.subplot_1.plot(self.x_cor_1, self.y_cor_1_1, color='r', label="object 1")
            self.subplot_1.plot(self.x_cor_1, self.y_cor_2_1, color='b', label="object 2")
            self.subplot_1.set_xlabel("Frame")
            self.subplot_1.set_ylabel("Position (mm)")
            # 给每条线添加标签（通过label参数）并显示图例（通过legend函数）
            self.subplot_1.legend(frameon=True)  # Show the legend
            self.inverse_relative_position_graph.tight_layout()
            self.inverse_relative_position_graph.subplots_adjust(top=0.9)
            self.canvas_1 = FigureCanvasTkAgg(self.inverse_relative_position_graph,
                                              master=self.inverse_relative_position_graph_frame)  # A tk.DrawingArea.
            self.canvas_1.draw()
            self.canvas_1.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
            self.line_1 = None
            self.annotation_1_1 = None
            self.annotation_2_1 = None
            self.canvas_1.mpl_connect('button_press_event', self.on_click)
            self.canvas_1.mpl_connect('motion_notify_event', self.on_move_1)
            self.frame_entry_container = tk.Frame(self.container_canvas_frame, width=int(half_width * 12 / 13),
                                                  height=int(toplevel_height * 0.6 / 5))
            self.frame_entry_container.pack()
            self.frame_label = tk.Label(self.frame_entry_container, text=f"Frame (1 - {total_frame}): ")
            self.frame_label.place(relx=0.4, rely=0.37, anchor="e")
            self.frame_entry = tk.Entry(self.frame_entry_container, textvariable=self.master.frame_enter_value)
            self.frame_entry.place(relx=0.42, rely=0.37, anchor="w")
            self.frame_btn = tk.Button(self.frame_entry_container, text="Submit", command=self.submit_frame_value)
            self.frame_btn.place(relx=0.5, rely=0.72, anchor="center")
            self.graph_2_frame = tk.Frame(self.container_canvas_frame, width=int(half_width * 12 / 13),
                                          height=int(toplevel_height * 3 / 5))
            self.graph_2_frame.pack()
            self.instantaneous_velocity_label = tk.Label(self.graph_2_frame, text="Instantaneous Velocity Graph (mm/s)")
            self.instantaneous_velocity_label.place(relx=0.5, rely=0.08, anchor='center')
            self.instantaneous_velocity_graph_frame = tk.Frame(self.graph_2_frame)
            self.instantaneous_velocity_graph_frame.place(relx=0.5, rely=0.55, relwidth=1, relheight=0.8,
                                                          anchor='center')
            self.instantaneous_velocity_graph = Figure()
            self.subplot_2 = self.instantaneous_velocity_graph.add_subplot(111)
            delta_y_list_1 = []
            for i in range(0, len(obj_1_cor) - 2, 2):
                delta_1 = obj_1_cor[i][1] - obj_1_cor[i + 2][1]
                delta_y_list_1.append(delta_1)
            delta_y_list_2 = []
            for j in range(0, len(obj_2_cor) - 2, 2):
                delta_2 = obj_2_cor[j][1] - obj_2_cor[j + 2][1]
                delta_y_list_2.append(delta_2)
            cap_2 = cv.VideoCapture(f'{self.master.processed_video_path_dynamic}/{self.master.basement_name_dynamic}')
            timestamp_list = []
            while True:
                ret, frame = cap_2.read()
                if not ret:
                    break
                # 获取当前帧的时间戳（秒）
                timestamp = cap_2.get(cv.CAP_PROP_POS_MSEC) / 1000
                timestamp_list.append(timestamp)
            delta_timestamp_list = []
            for k in range(0, len(timestamp_list) - 2, 2):
                delta_ts = timestamp_list[k + 2] - timestamp_list[k]
                delta_timestamp_list.append(delta_ts)
            self.x_cor_2 = [(frame + 1) for frame in range(0, total_frame, 2)]
            self.y_cor_1_2 = [0]
            self.y_cor_2_2 = [0]
            for i in range(len(delta_timestamp_list)):
                ins_vel_1 = ((delta_y_list_1[i] / height_ratio) * 10) / delta_timestamp_list[i]
                self.y_cor_1_2.append(ins_vel_1)
                ins_vel_2 = ((delta_y_list_2[i] / height_ratio) * 10) / delta_timestamp_list[i]
                self.y_cor_2_2.append(ins_vel_2)
            self.subplot_2.plot(self.x_cor_2, self.y_cor_1_2, color='r', label="object 1")
            self.subplot_2.plot(self.x_cor_2, self.y_cor_2_2, color='b', label="object 2")
            self.threshold = 1
            self.subplot_2.set_xlabel("Frame")
            self.subplot_2.set_ylabel("Velocity (mm/s)")
            # 给每条线添加标签（通过label参数）并显示图例（通过legend函数）
            self.subplot_2.legend(frameon=True)  # Show the legend
            self.instantaneous_velocity_graph.tight_layout()
            self.instantaneous_velocity_graph.subplots_adjust(top=0.9)
            self.canvas_2 = FigureCanvasTkAgg(self.instantaneous_velocity_graph,
                                              master=self.instantaneous_velocity_graph_frame)  # A tk.DrawingArea.
            self.canvas_2.draw()
            self.canvas_2.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
            self.line_2 = None
            self.annotation_1_2 = None
            self.annotation_2_2 = None
            self.canvas_2.mpl_connect('button_press_event', self.on_click)
            self.canvas_2.mpl_connect('motion_notify_event', self.on_move_2)
            self.metrics_container = tk.Frame(self.container_canvas_frame, width=int(half_width * (12 / 13)),
                                              height=toplevel_height)
            self.metrics_container.pack()
            self.T_label = tk.Label(self.metrics_container,
                                    text=f"T is the number of processed video frames, T = {total_frame}")
            self.T_label.pack(pady=10)
            self.prox_frame = tk.Frame(self.metrics_container, width=int(half_width * (12 / 13)),
                                       height=toplevel_height * 1 / 3)
            self.prox_frame.pack()
            self.prox_result_label = tk.Label(self.prox_frame, text="1. Proximity index (Prox)")
            self.prox_result_label.place(relx=0.02, rely=0.04, anchor="nw")
            self.prox_result_treeview_frame = tk.Frame(self.prox_frame)
            self.prox_result_treeview_frame.place(relx=0.5, rely=0.55, relwidth=1, relheight=0.8, anchor='center')
            self.prox_result_treeview = ttk.Treeview(self.prox_result_treeview_frame, columns=("c1", "c2", "c3", "c4"),
                                                     show='headings')
            self.prox_result_treeview.column("c1", anchor=tk.CENTER)
            self.prox_result_treeview.heading("c1", text="Author",
                                              command=lambda _col="c1": self.sortby(_col, 0, self.prox_result_treeview))
            self.prox_result_treeview.column("c2", anchor=tk.CENTER)
            self.prox_result_treeview.heading("c2", text="δ",
                                              command=lambda _col="c2": self.sortby(_col, 0, self.prox_result_treeview))
            self.prox_result_treeview.column("c3", anchor=tk.CENTER)
            self.prox_result_treeview.heading("c3", text="Prox",
                                              command=lambda _col="c3": self.sortby(_col, 0, self.prox_result_treeview))
            self.prox_result_treeview.column("c4", width=0, minwidth=0, stretch='no')
            self.prox_result_treeview.heading("c4", text="Timestamp",
                                              command=lambda _col="c4": self.sortby(_col, 0, self.prox_result_treeview))
            self.prox_result_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            # Create a vertical scrollbar
            self.prox_result_treeview_vsb = ttk.Scrollbar(self.prox_result_treeview_frame, orient="vertical",
                                                          command=self.prox_result_treeview.yview)
            self.prox_result_treeview_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.prox_result_treeview.configure(yscrollcommand=self.prox_result_treeview_vsb.set)
            # prox pre-insert
            for prox in data["prox"]:
                self.prox_result_treeview.insert('', 'end',
                                                 values=(prox["author"], prox["δ"], prox["prox"], prox["timestamp"]))
            self.csem_frame = tk.Frame(self.metrics_container, width=int(half_width * (12 / 13)),
                                       height=toplevel_height * 1 / 3)
            self.csem_frame.pack()
            self.csem_result_label = tk.Label(self.csem_frame, text="2. Cross sampled entropy (CSEM)")
            self.csem_result_label.place(relx=0.02, rely=0.04, anchor="nw")
            self.csem_result_treeview_frame = tk.Frame(self.csem_frame)
            self.csem_result_treeview_frame.place(relx=0.5, rely=0.55, relwidth=1, relheight=0.8, anchor='center')
            self.csem_result_treeview = ttk.Treeview(self.csem_result_treeview_frame,
                                                     columns=("c1", "c2", "c3", "c4", "c5"), show='headings')
            self.csem_result_treeview.column("c1", anchor=tk.CENTER, width=int(half_width * (1 / 6)))
            self.csem_result_treeview.heading("c1", text="Author",
                                              command=lambda _col="c1": self.sortby(_col, 0, self.csem_result_treeview))
            self.csem_result_treeview.column("c2", anchor=tk.CENTER, width=int(half_width * (1 / 6)))
            self.csem_result_treeview.heading("c2", text="δ",
                                              command=lambda _col="c2": self.sortby(_col, 0, self.csem_result_treeview))
            self.csem_result_treeview.column("c3", anchor=tk.CENTER, width=int(half_width * (1 / 6)))
            self.csem_result_treeview.heading("c3", text="max(m)",
                                              command=lambda _col="c3": self.sortby(_col, 0, self.csem_result_treeview))
            self.csem_result_treeview.column("c4", anchor=tk.CENTER, width=int(half_width * (1 / 6)))
            self.csem_result_treeview.heading("c4", text="CSEM",
                                              command=lambda _col="c4": self.sortby(_col, 0, self.csem_result_treeview))
            self.csem_result_treeview.column("c5", width=0, minwidth=0, stretch='no')
            self.csem_result_treeview.heading("c5", text="Timestamp",
                                              command=lambda _col="c5": self.sortby(_col, 0, self.csem_result_treeview))
            self.csem_result_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            # Create a vertical scrollbar
            self.csem_result_treeview_vsb = ttk.Scrollbar(self.csem_result_treeview_frame, orient="vertical",
                                                          command=self.csem_result_treeview.yview)
            self.csem_result_treeview_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.csem_result_treeview.configure(yscrollcommand=self.csem_result_treeview_vsb.set)
            for csem in data["csem"]:
                self.csem_result_treeview.insert('', 'end', values=(
                csem["author"], csem["δ"], csem["max(m)"], csem["CSEM"], csem["timestamp"]))
            self.cs_frame = tk.Frame(self.metrics_container, width=int(half_width * (12 / 13)),
                                     height=toplevel_height * 1 / 3)
            self.cs_frame.pack()
            self.cs_result_label = tk.Label(self.cs_frame, text="3. Coefficient of Sociality (Cs)")
            self.cs_result_label.place(relx=0.02, rely=0.04, anchor="nw")
            if data["cs"] != []:
                self.cs_result_DO_label = tk.Label(self.cs_frame, text=f'D_O = {data["cs"][0]["D_O"]}')
                self.cs_result_DO_label.place(relx=0.5, rely=0.3, anchor="center")
                self.cs_result_DE_label = tk.Label(self.cs_frame, text=f'D_E = {data["cs"][0]["D_E"]}')
                self.cs_result_DE_label.place(relx=0.5, rely=0.5, anchor="center")
                self.cs_result_Cs_label = tk.Label(self.cs_frame, text=f'Cs = {data["cs"][0]["Cs"]}')
                self.cs_result_Cs_label.place(relx=0.5, rely=0.7, anchor="center")
            else:
                distance_list = data["distancelist"]
                object_1_coordinate = data['object1coordinatelist']
                object_2_coordinate = data['object2coordinatelist']
                all_distance_list = []
                for cor1 in object_1_coordinate:
                    for cor2 in object_2_coordinate:
                        distance = round(
                            math.sqrt(
                                ((cor1[0] - cor2[0]) / width_ratio) ** 2 + ((cor1[1] - cor2[1]) / height_ratio) ** 2),
                            2)
                        all_distance_list.append(distance)
                total_frame_count = len(distance_list)
                # cs DO value
                DO_value = sum(distance_list) / total_frame_count
                self.cs_result_DO_label = tk.Label(self.cs_frame, text=f"D_O = {DO_value}")
                self.cs_result_DO_label.place(relx=0.5, rely=0.3, anchor="center")
                # cs DE value
                DE_value = sum(all_distance_list) / total_frame_count ** 2
                self.cs_result_DE_label = tk.Label(self.cs_frame, text=f"D_E = {DE_value}")
                self.cs_result_DE_label.place(relx=0.5, rely=0.5, anchor="center")
                # cs Cs value
                Cs_value = (DE_value - DO_value) / (DE_value + DO_value)
                self.cs_result_Cs_label = tk.Label(self.cs_frame, text=f"Cs = {Cs_value}")
                self.cs_result_Cs_label.place(relx=0.5, rely=0.7, anchor="center")
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                    # 对数据进行更新
                    data['cs'] = [{
                        "D_O": DO_value,
                        "D_E": DE_value,
                        "Cs": Cs_value
                    }]
                    # 将更新后的数据写回文件
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)
                # 写入word文档
                # 创建Cs word文档并编辑写入Cs result
                # Add prox to metrics file as .docx
                frame_cs_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Cs.docx'
                # os.path.exists()是Python中os模块的一个方法，用于检查指定的路径（文件或目录）是否存在。
                # 如果路径存在，它会返回True；如果路径不存在，它会返回False。
                if os.path.exists(frame_cs_docx_path):
                    doc = Document(frame_cs_docx_path)
                else:
                    doc = Document()  # 创建一个Document对象
                # 遍历文档中的所有段落，并删除
                for element in doc.element.body:
                    doc.element.body.remove(element)
                doc.add_paragraph(
                    f'D_O = {DO_value}\nD_E = {DE_value}\nCs = {Cs_value}')  # 在文档中添加一个段落
                doc.save(frame_cs_docx_path)  # 将文档保存为一个.docx文件
            self.rv_result_frame = tk.Frame(self.metrics_container, width=int(half_width * (12 / 13)),
                                            height=toplevel_height * 1 / 3)
            self.rv_result_frame.pack()
            self.rv_result_title_label = tk.Label(self.rv_result_frame, text="4. Correlation Coefficient (r_V)")
            self.rv_result_title_label.place(relx=0.02, rely=0.04, anchor="nw")
            if data["rv"] != []:
                self.rv_result_latitude_label = tk.Label(self.rv_result_frame,
                                                         text=f'r_Latitude = {data["rv"][0]["r_Latitude"]}')
                self.rv_result_latitude_label.place(relx=0.5, rely=0.33, anchor="center")
                self.rv_result_speed_label = tk.Label(self.rv_result_frame,
                                                      text=f'r_Speed = {data["rv"][0]["r_Speed"]}')
                self.rv_result_speed_label.place(relx=0.5, rely=0.66, anchor="center")
            else:
                object_1_lat_list = []
                for coordinate_5 in obj_1_cor:
                    object_1_lat_list.append(coordinate_5[1])
                avg_lat_object_1 = np.mean(object_1_lat_list)
                object_2_lat_list = []
                for coordinate_6 in obj_2_cor:
                    object_2_lat_list.append(coordinate_6[1])
                avg_lat_object_2 = np.mean(object_2_lat_list)
                # 分子部分：求和((x_i - mean_x) * (y_i - mean_y))的累加值
                numerator_2 = sum(
                    (object_1_lat_list - avg_lat_object_1) * (object_2_lat_list - avg_lat_object_2))
                # 分母部分：开方(求和((x_i - mean_x) ^ 2) * 求和((y_i - mean_y) ^ 2))
                denominator_2 = np.sqrt(sum((object_1_lat_list - avg_lat_object_1) ** 2) * sum(
                    (object_2_lat_list - avg_lat_object_2) ** 2))
                # 计算皮尔逊相关系数
                pearson_coefficient_2 = numerator_2 / denominator_2
                speed_list_1 = []
                for i in range(1, len(object_1_lat_list)):
                    # 速度是位置的变化除以时间的变化（在这里时间的变化总是1帧）
                    speed_1 = abs(object_1_lat_list[i] - object_1_lat_list[i - 1])
                    speed_list_1.append(speed_1)
                avg_speed_1 = np.mean(speed_list_1)
                speed_list_2 = []
                for j in range(1, len(object_2_lat_list)):
                    # 速度是位置的变化除以时间的变化（在这里时间的变化总是1秒）
                    speed_2 = abs(object_2_lat_list[j] - object_2_lat_list[j - 1])
                    speed_list_2.append(speed_2)
                avg_speed_2 = np.mean(speed_list_2)
                # 分子部分：求和((x_i - mean_x) * (y_i - mean_y))的累加值
                numerator_4 = sum(
                    (speed_list_1 - avg_speed_1) * (speed_list_2 - avg_speed_2))
                # 分母部分：开方(求和((x_i - mean_x) ^ 2) * 求和((y_i - mean_y) ^ 2))
                denominator_4 = np.sqrt(sum((speed_list_1 - avg_speed_1) ** 2) * sum(
                    (speed_list_2 - avg_speed_2) ** 2))
                # 计算皮尔逊相关系数
                pearson_coefficient_4 = numerator_4 / denominator_4
                self.rv_result_latitude_label = tk.Label(self.rv_result_frame,
                                                         text=f'r_Latitude = {pearson_coefficient_2}')
                self.rv_result_latitude_label.place(relx=0.5, rely=0.33, anchor="center")
                self.rv_result_speed_label = tk.Label(self.rv_result_frame,
                                                      text=f'r_Speed = {pearson_coefficient_4}')
                self.rv_result_speed_label.place(relx=0.5, rely=0.66, anchor="center")
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                    # 对数据进行更新
                    data['rv'] = [{
                        "r_Latitude": pearson_coefficient_2,
                        "r_Speed": pearson_coefficient_4
                    }]
                    # 将更新后的数据写回文件
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)
                # 写入word文档
                frame_rv_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Correlation Coefficient.docx'
                # os.path.exists()是Python中os模块的一个方法，用于检查指定的路径（文件或目录）是否存在。
                # 如果路径存在，它会返回True；如果路径不存在，它会返回False。
                if os.path.exists(frame_rv_docx_path):
                    doc = Document(frame_rv_docx_path)
                else:
                    doc = Document()  # 创建一个Document对象
                # 遍历文档中的所有段落，并删除
                for element in doc.element.body:
                    doc.element.body.remove(element)
                doc.add_paragraph(
                    f'r_Latitude = {pearson_coefficient_2}\nr_Speed = {pearson_coefficient_4}')  # 在文档中添加一个段落
                doc.save(frame_rv_docx_path)  # 将文档保存为一个.docx文件
            self.di_result_frame = tk.Frame(self.metrics_container, width=int(half_width * (12 / 13)),
                                            height=toplevel_height * 1 / 3)
            self.di_result_frame.pack()
            self.di_result_title_label = tk.Label(self.di_result_frame, text="5. Dynamic Interaction (DI)")
            self.di_result_title_label.place(relx=0.02, rely=0.04, anchor="nw")
            if data["di"] != []:
                self.di_result_dis_label = tk.Label(self.di_result_frame,
                                                    text=f'DI_d = {data["di"][0]["DI_d"]}')
                self.di_result_dis_label.place(relx=0.5, rely=0.25, anchor="center")
                self.di_result_angle_label = tk.Label(self.di_result_frame,
                                                      text=f'DI_θ = {data["di"][0]["DI_θ"]}')
                self.di_result_angle_label.place(relx=0.5, rely=0.5, anchor="center")
                self.di_result_com_label = tk.Label(self.di_result_frame,
                                                    text=f'DI = {data["di"][0]["DI"]}')
                self.di_result_com_label.place(relx=0.5, rely=0.75, anchor="center")
            else:
                object_1_position_list = []
                for coordinate_7 in obj_1_cor:
                    object_1_position_list.append(coordinate_7[1])
                object_2_position_list = []
                for coordinate_8 in obj_2_cor:
                    object_2_position_list.append(coordinate_8[1])
                object_1_displacement_list = [abs(object_1_position_list[i + 1] - object_1_position_list[i]) for i
                                              in range(len(object_1_position_list) - 1)]
                object_2_displacement_list = [abs(object_2_position_list[i + 1] - object_2_position_list[i]) for i
                                              in range(len(object_2_position_list) - 1)]
                T_variable = len(object_1_displacement_list)
                total_sum_di_dis = 0
                for t in range(T_variable):
                    denominator = object_1_displacement_list[t] + object_2_displacement_list[t]
                    if denominator != 0:
                        total_sum_di_dis += (1 - abs(
                            object_1_displacement_list[t] - object_2_displacement_list[t]) / denominator)
                    else:
                        total_sum_di_dis += 1
                DI_d = total_sum_di_dis / T_variable if T_variable != 0 else 0  # 避免除以0
                object_1_displacement_neg_list = [object_1_position_list[i + 1] - object_1_position_list[i] for i in
                                                  range(len(object_1_position_list) - 1)]
                object_2_displacement_neg_list = [object_2_position_list[i + 1] - object_2_position_list[i] for i in
                                                  range(len(object_2_position_list) - 1)]
                object_1_angle_list = [90 if y < 0 else (-90 if y > 0 else 0) for y in object_1_displacement_neg_list]
                object_2_angle_list = [90 if y < 0 else (-90 if y > 0 else 0) for y in object_2_displacement_neg_list]
                theta_diff_list = [a - b for a, b in zip(object_1_angle_list, object_2_angle_list)]
                total_sum_di_theta = 0
                for t in range(T_variable):
                    absolute_angle = round(math.cos(math.radians(theta_diff_list[t])), 2)
                    total_sum_di_theta += absolute_angle
                DI_theta = total_sum_di_theta / T_variable
                total_sum_di = 0
                for t in range(T_variable):
                    # 位移差的和
                    displacement_sum = object_1_displacement_list[t] + object_2_displacement_list[t]
                    term1 = round(math.cos(math.radians(theta_diff_list[t])), 2)
                    # 排除除数为0的情况
                    if displacement_sum != 0:
                        term2 = 1 - (abs(
                            object_1_displacement_list[t] - object_2_displacement_list[t])) / displacement_sum
                    else:
                        term2 = 1
                    total_sum_di += term1 * term2
                DI_com = total_sum_di / T_variable if T_variable != 0 else 0  # 避免除以0
                self.di_result_dis_label = tk.Label(self.di_result_frame,
                                                    text=f'DI_d = {DI_d}')
                self.di_result_dis_label.place(relx=0.5, rely=0.25, anchor="center")
                self.di_result_angle_label = tk.Label(self.di_result_frame,
                                                      text=f'DI_θ = {DI_theta}')
                self.di_result_angle_label.place(relx=0.5, rely=0.5, anchor="center")
                self.di_result_com_label = tk.Label(self.di_result_frame,
                                                    text=f'DI = {DI_com}')
                self.di_result_com_label.place(relx=0.5, rely=0.75, anchor="center")
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)
                    # 对数据进行更新
                    data['di'] = [{
                        "DI_d": DI_d,
                        "DI_θ": DI_theta,
                        "DI": DI_com
                    }]
                    # 将更新后的数据写回文件
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)
                # 写入word文档
                frame_di_docx_path = f'{self.master.experiment_path_dynamic}/Metrics_Result/Dynamic Interaction.docx'
                # os.path.exists()是Python中os模块的一个方法，用于检查指定的路径（文件或目录）是否存在。
                # 如果路径存在，它会返回True；如果路径不存在，它会返回False。
                if os.path.exists(frame_di_docx_path):
                    doc = Document(frame_di_docx_path)
                else:
                    doc = Document()  # 创建一个Document对象
                # 遍历文档中的所有段落，并删除
                for element in doc.element.body:
                    doc.element.body.remove(element)
                doc.add_paragraph(
                    f'DI_d = {DI_d}\nDI_θ = {DI_theta}\nDI = {DI_com}')  # 在文档中添加一个段落
                doc.save(frame_di_docx_path)  # 将文档保存为一个.docx文件

            # after方法会在100毫秒后调用move_scrollbar_to_top函数。
            def move_scrollbar_to_top():
                self.canvas.yview_moveto(0)

            self.result_toplevel.after(100, move_scrollbar_to_top)
            self.master.protocol("WM_DELETE_WINDOW", self.on_closing)

        def submit_frame_value(self):
            def is_integer(entry):
                try:
                    num = int(entry)
                except ValueError:
                    return False
                return True

            if is_integer(self.master.frame_enter_value.get().strip()) == True:
                if int(self.master.frame_enter_value.get().strip()) >= 1 and int(
                        self.master.frame_enter_value.get().strip()) <= self.total_frame:
                    # 遍历 Treeview 中的每个项目
                    wait_selected_id_list = []
                    for item in self.annotation_all_tree.get_children():
                        # 如果找到了与输入值相符的 Frame 值，那么选择该项目
                        if int(self.annotation_all_tree.item(item, "values")[0]) == int(
                                self.master.frame_enter_value.get().strip()):
                            wait_selected_id_list.append(item)
                    self.annotation_all_tree.selection_set(wait_selected_id_list)
                    if wait_selected_id_list:
                        self.annotation_all_tree.see(wait_selected_id_list[0])
                        full_content = self.all_full_contents[wait_selected_id_list[0]]
                        self.display_all_text.delete(1.0, tk.END)
                        self.display_all_text.insert(tk.END, full_content)
                    # Set video frame position based on selected frame
                    self.cap_1.set(cv.CAP_PROP_POS_FRAMES, int(self.master.frame_enter_value.get().strip()) - 1)
                    # Set slider position based on selected frame
                    self.slider_1.set(int(self.master.frame_enter_value.get().strip()))
                    if self.line_1:
                        self.line_1.remove()
                    if self.line_2:
                        self.line_2.remove()
                    if self.annotation_1_1:
                        self.annotation_1_1.remove()
                    if self.annotation_2_1:
                        self.annotation_2_1.remove()
                    if self.annotation_1_2:
                        self.annotation_1_2.remove()
                    if self.annotation_2_2:
                        self.annotation_2_2.remove()
                    self.line_1 = self.subplot_1.axvline(x=int(self.master.frame_enter_value.get().strip()), color='g')
                    self.line_2 = self.subplot_2.axvline(x=int(self.master.frame_enter_value.get().strip()), color='g')
                    y_interp_1_1 = np.interp(int(self.master.frame_enter_value.get().strip()), self.x_cor_1,
                                             self.y_cor_1_1)
                    y_interp_2_1 = np.interp(int(self.master.frame_enter_value.get().strip()), self.x_cor_1,
                                             self.y_cor_2_1)
                    y_interp_1_2 = np.interp(int(self.master.frame_enter_value.get().strip()), self.x_cor_2,
                                             self.y_cor_1_2)
                    y_interp_2_2 = np.interp(int(self.master.frame_enter_value.get().strip()), self.x_cor_2,
                                             self.y_cor_2_2)

                    # 显示坐标
                    self.annotation_1_1 = self.subplot_1.annotate(
                        f'({int(self.master.frame_enter_value.get().strip()):.2f}, {y_interp_1_1:.2f})',
                        xy=(int(self.master.frame_enter_value.get().strip()), y_interp_1_1),
                        xytext=(int(self.master.frame_enter_value.get().strip()), y_interp_1_1))
                    self.annotation_2_1 = self.subplot_1.annotate(
                        f'({int(self.master.frame_enter_value.get().strip()):.2f}, {y_interp_2_1:.2f})',
                        xy=(int(self.master.frame_enter_value.get().strip()), y_interp_2_1),
                        xytext=(int(self.master.frame_enter_value.get().strip()), y_interp_2_1))
                    self.annotation_1_2 = self.subplot_2.annotate(
                        f'({int(self.master.frame_enter_value.get().strip()):.2f}, {y_interp_1_2:.2f})',
                        xy=(int(self.master.frame_enter_value.get().strip()), y_interp_1_2),
                        xytext=(int(self.master.frame_enter_value.get().strip()), y_interp_1_2))
                    self.annotation_2_2 = self.subplot_2.annotate(
                        f'({int(self.master.frame_enter_value.get().strip()):.2f}, {y_interp_2_2:.2f})',
                        xy=(int(self.master.frame_enter_value.get().strip()), y_interp_2_2),
                        xytext=(int(self.master.frame_enter_value.get().strip()), y_interp_2_2))

                    self.canvas_1.draw()
                    self.canvas_2.draw()

                else:
                    messagebox.showwarning("Warning",
                                           f"Please enter an integer within the range from 1 to {self.total_frame}.")
            else:
                messagebox.showwarning("Warning",
                                       f"Please enter an integer within the range from 1 to {self.total_frame}.")

        def on_click(self, event):
            x = event.xdata
            # 检查了鼠标点击事件是否在图像的坐标轴内
            if event.inaxes is not None:
                if self.line_1 in self.subplot_1.lines:
                    self.line_1.remove()
                if self.line_2 in self.subplot_2.lines:
                    self.line_2.remove()
                if self.annotation_1_1 in self.subplot_1.texts:
                    self.annotation_1_1.remove()
                if self.annotation_2_1 in self.subplot_1.texts:
                    self.annotation_2_1.remove()
                if self.annotation_1_2 in self.subplot_2.texts:
                    self.annotation_1_2.remove()
                if self.annotation_2_2 in self.subplot_2.texts:
                    self.annotation_2_2.remove()

                self.line_1 = self.subplot_1.axvline(x=x, color='g')
                self.line_2 = self.subplot_2.axvline(x=x, color='g')

                y_interp_1_1 = np.interp(x, self.x_cor_1, self.y_cor_1_1)
                y_interp_2_1 = np.interp(x, self.x_cor_1, self.y_cor_2_1)
                y_interp_1_2 = np.interp(x, self.x_cor_2, self.y_cor_1_2)
                y_interp_2_2 = np.interp(x, self.x_cor_2, self.y_cor_2_2)

                # 显示坐标
                self.annotation_1_1 = self.subplot_1.annotate(f'({x:.2f}, {y_interp_1_1:.2f})', xy=(x, y_interp_1_1),
                                                              xytext=(x, y_interp_1_1))
                self.annotation_2_1 = self.subplot_1.annotate(f'({x:.2f}, {y_interp_2_1:.2f})', xy=(x, y_interp_2_1),
                                                              xytext=(x, y_interp_2_1))
                self.annotation_1_2 = self.subplot_2.annotate(f'({x:.2f}, {y_interp_1_1:.2f})', xy=(x, y_interp_1_2),
                                                              xytext=(x, y_interp_1_2))
                self.annotation_2_2 = self.subplot_2.annotate(f'({x:.2f}, {y_interp_2_1:.2f})', xy=(x, y_interp_2_2),
                                                              xytext=(x, y_interp_2_2))

                self.canvas_1.draw()
                self.canvas_2.draw()

                frame = int(x)
                self.frame_entry.delete(0, tk.END)  # 删除现有的内容
                self.frame_entry.insert(0, f'{frame}')  # 插入新的内容
                # Set video frame position based on selected frame
                # Note that we subtract 1 from the total_frames because frame indices are 0-based.
                self.cap_1.set(cv.CAP_PROP_POS_FRAMES, frame - 1)
                # Set slider position based on selected frame
                self.slider_1.set(frame)
                # item_id is a list
                item_id_list = self.annotation_all_tree.get_children()
                item_id_selection_list = []
                for i, item_id in enumerate(item_id_list):
                    # {'text': '', 'image': '', 'values': [1, 'Yuan', 'make it mind', '2023-07-18T13:04:12.255241'], 'open': 0, 'tags': ''}
                    item = self.annotation_all_tree.item(item_id)  # Get the item
                    frame = item['values'][0]
                    if int(x) == frame:
                        item_id_selection_list.append(item_id_list[i])
                if item_id_selection_list:
                    self.annotation_all_tree.selection_set(item_id_selection_list)
                    self.update_sort(self.annotation_all_tree)
                    self.annotation_all_tree.see(item_id_selection_list[0])
                    full_content = self.all_full_contents[item_id_selection_list[0]]
                    self.display_all_text.delete(1.0, tk.END)
                    self.display_all_text.insert(tk.END, full_content)
                else:
                    self.annotation_all_tree.selection_remove(self.annotation_all_tree.selection())
                    self.update_sort(self.annotation_all_tree)
                    self.display_all_text.delete(1.0, tk.END)

        def on_move_1(self, event):
            if event.inaxes is not None:
                x = event.xdata
                y = event.ydata
                self.subplot_1.set_title(f'x={x:.2f}, y={y:.2f}')
                self.canvas_1.draw()
            else:
                self.subplot_1.set_title("")
                self.canvas_1.draw()

        def on_move_2(self, event):
            if event.inaxes is not None:
                x = event.xdata
                y = event.ydata
                self.subplot_2.set_title(f'x={x:.2f}, y={y:.2f}')
                self.canvas_2.draw()
            else:
                self.subplot_2.set_title("")
                self.canvas_2.draw()

        def on_all_treeview_select(self, event):
            item_id = self.annotation_all_tree.selection()
            if len(item_id) == 1:
                item = self.annotation_all_tree.item(item_id)  # Get the item
                frame = item['values'][0]  # The frame value is the first value
                full_content = self.all_full_contents[item_id[0]]
                self.display_all_text.delete(1.0, tk.END)
                self.display_all_text.insert(tk.END, full_content)
                # Set video frame position based on selected frame
                self.cap_1.set(cv.CAP_PROP_POS_FRAMES, frame - 1)
                # Set slider position based on selected frame
                self.slider_1.set(frame)
                x = frame
                if self.line_1:
                    self.line_1.remove()
                if self.line_2:
                    self.line_2.remove()
                self.line_1 = self.subplot_1.axvline(x=x, color='g')
                self.line_2 = self.subplot_2.axvline(x=x, color='g')
                self.canvas_1.draw()
                self.canvas_2.draw()

        def slider_moved_1(self, event):
            # Only update frame when video is paused
            if self.pause_video_1:
                # Change frame position based on slider
                self.cap_1.set(cv.CAP_PROP_POS_FRAMES, self.slider_1.get() - 1)
                # Play one frame
                success, frame = self.cap_1.read()
                if success:
                    height, width, _ = frame.shape
                    aspect_ratio = width / height
                    if width > self.screen_width or height > self.screen_height:
                        if self.screen_width / aspect_ratio <= self.screen_height:
                            new_width = self.screen_width
                            new_height = int(self.screen_width / aspect_ratio)
                        else:
                            new_height = self.screen_height
                            new_width = int(self.screen_height * aspect_ratio)
                        new_height = int(new_height * 3.2 / 5)
                        new_width = int(new_width * 3.2 / 5)
                    else:
                        new_height = int(height * 3.2 / 5)
                        new_width = int(width * 3.2 / 5)
                    # self.master.scale_length_1 = new_width
                    img = cv.cvtColor(frame, cv.COLOR_BGR2RGBA)
                    current_image = PIL.Image.fromarray(img).resize((new_width, new_height), PIL.Image.LANCZOS)
                    enhancer = ImageEnhance.Sharpness(current_image)
                    current_image = enhancer.enhance(2.0)  # Increase sharpness
                    # current_image = PIL.Image.fromarray(img).resize((500, 500))
                    imgtk = PIL.ImageTk.PhotoImage(image=current_image)
                    self.movieLabel_1.imgtk = imgtk
                    self.movieLabel_1.config(image=imgtk)
                    self.result_toplevel.update()
            else:
                self.cap_1.set(cv.CAP_PROP_POS_FRAMES, self.slider_1.get() - 1)

        def video_player_1(self):
            while self.master.thread_running:
                read_first_frame = False
                while self.cap_1.isOpened():
                    if self.pause_video_1 and read_first_frame:
                        continue
                    success, frame = self.cap_1.read()
                    read_first_frame = True
                    if success == True:
                        height, width, _ = frame.shape
                        aspect_ratio = width / height
                        if width > self.screen_width or height > self.screen_height:
                            if self.screen_width / aspect_ratio <= self.screen_height:
                                new_width = self.screen_width
                                new_height = int(self.screen_width / aspect_ratio)
                            else:
                                new_height = self.screen_height
                                new_width = int(self.screen_height * aspect_ratio)
                            new_height = int(new_height * 3.2 / 5)
                            new_width = int(new_width * 3.2 / 5)
                        else:
                            new_height = int(height * 3.2 / 5)
                            new_width = int(width * 3.2 / 5)
                        # self.master.scale_length_1 = new_width
                        # 使Label框中显示指定帧
                        img = cv.cvtColor(frame, cv.COLOR_BGR2RGBA)
                        current_image = PIL.Image.fromarray(img).resize((new_width, new_height), PIL.Image.LANCZOS)
                        enhancer = ImageEnhance.Sharpness(current_image)
                        current_image = enhancer.enhance(2.0)  # Increase sharpness
                        imgtk = PIL.ImageTk.PhotoImage(image=current_image)
                        try:
                            self.movieLabel_1.imgtk = imgtk
                            if self.movieLabel_1.winfo_exists():
                                self.movieLabel_1.config(image=imgtk)
                                self.slider_1.set(self.cap_1.get(cv.CAP_PROP_POS_FRAMES))
                                self.result_toplevel.update()
                        except:
                            continue
                    else:
                        self.pause_video_1 = True
                        try:
                            self.play_pause_btn_1.config(text='Play')
                        except:
                            continue

                self.cap_1.release()
                time.sleep(1)

        def play_pause_1(self):  # toggle play/pause status
            self.pause_video_1 = not self.pause_video_1
            if self.pause_video_1:
                self.play_pause_btn_1.config(text='Play')
            else:
                self.play_pause_btn_1.config(text='Pause')

        def back_to_upload_page_or_experiment_page(self):
            from UploadPage import UploadPage
            from ExperimentPage import ExperimentPage
            directory_path = os.path.abspath(os.getcwd())
            self.master.repository_path_dynamic = f'{directory_path}/Repository'
            if self.master.selected_exp == "":
                self.master.selected_new = ""
                self.master.show_frame(UploadPage)
            else:
                self.master.show_frame(ExperimentPage)
                self.master.selected_exp == ""

        def manage_own_annotations(self):
            # Delete all children of self.right_frame
            for widget in self.right_frame.winfo_children():
                widget.destroy()

            # 在右侧Frame中创建一个内嵌的Frame来装Own Treeview和Own Scrollbar
            self.annotation_own_tree_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
            self.annotation_own_tree_frame.place(relx=0.5, rely=0.34, relwidth=0.9, relheight=0.45, anchor='center')
            self.annotation_own_tree = ttk.Treeview(self.annotation_own_tree_frame, columns=("c1", "c2", "c3", "c4"),
                                                    show='headings', selectmode="extended")
            self.annotation_own_tree.column("c1", anchor=tk.CENTER, width=25)
            self.annotation_own_tree.heading("c1", text="Frame",
                                             command=lambda _col="c1": self.sortby(_col, 0, self.annotation_own_tree))
            self.annotation_own_tree.column("c2", anchor=tk.CENTER, width=50)
            self.annotation_own_tree.heading("c2", text="Author",
                                             command=lambda _col="c2": self.sortby(_col, 0, self.annotation_own_tree))
            self.annotation_own_tree.column("c3", anchor=tk.CENTER, width=325)
            self.annotation_own_tree.heading("c3", text="Notes",
                                             command=lambda _col="c3": self.sortby(_col, 0, self.annotation_own_tree))
            self.annotation_own_tree.column("c4", width=0, minwidth=0, stretch='no')
            self.annotation_own_tree.heading("c4", text="Timestamp",
                                             command=lambda _col="c4": self.sortby(_col, 0, self.annotation_own_tree))
            self.annotation_own_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            self.annotation_own_tree.bind('<<TreeviewSelect>>', self.on_own_treeview_select)
            # 创建一个字典用来保存每个 Treeview 项目的完整内容
            self.own_full_contents = {}
            # Create a vertical scrollbar
            self.annotation_own_tree_vsb = ttk.Scrollbar(self.annotation_own_tree_frame, orient="vertical",
                                                         command=self.annotation_own_tree.yview)
            self.annotation_own_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
            # Configure the treeview to use the scrollbar
            self.annotation_own_tree.configure(yscrollcommand=self.annotation_own_tree_vsb.set)

            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            for note in data["annotation"]:
                current_frame = note["frame"]
                annotation_author = note["author"]
                text_content = note["text"]
                timestamp = note["timestamp"]
                split_text_content = text_content.split("\n")[0] if "\n" in text_content else text_content
                # Add annotation to treeview
                if annotation_author == self.master.username.get().strip():
                    item_id = self.annotation_own_tree.insert('', 'end', values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.own_full_contents[item_id] = text_content

            # 创建容纳Text的Frame
            self.text_own_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
            self.text_own_frame.place(relx=0.495, rely=0.75, relwidth=0.9, relheight=0.32, anchor='center')
            # 创建Text
            self.display_own_text = tk.Text(self.text_own_frame, width=40)
            self.display_own_text.pack(fill=tk.BOTH, expand=True)

            self.annotation_own_label = tk.Label(self.right_frame,
                                                 text=f'{self.master.username.get().strip()}, You can edit your own annotations.')
            self.annotation_own_label.place(relx=0.5, rely=0.08, anchor='center')

            self.back_button = tk.Button(self.right_frame, text="Back", command=self.back_to_initial, cursor="hand2")
            self.back_button.place(relx=0.1, rely=0.04, anchor='center')

            self.edit_own_btn = tk.Button(self.right_frame, text="Edit", command=self.edit_own_annotation,
                                          cursor="hand2")
            self.edit_own_btn.place(relx=0.1, rely=0.95, anchor='center')

            self.delete_own_btn = tk.Button(self.right_frame, text="Delete", command=self.delete_own_annotation,
                                            cursor="hand2")
            self.delete_own_btn.place(relx=0.88, rely=0.95, anchor='center')

        def on_own_treeview_select(self, event):
            item_id = self.annotation_own_tree.focus()
            item = self.annotation_own_tree.item(item_id)  # Get the item
            if not item['values']:  # Check if 'values' is empty
                return  # If empty, return immediately
            frame = item['values'][0]  # The frame value is the first value
            full_content = self.own_full_contents[item_id]
            self.display_own_text.delete(1.0, tk.END)
            self.display_own_text.insert(tk.END, full_content)
            # Set video frame position based on selected frame
            self.cap.set(cv.CAP_PROP_POS_FRAMES, frame - 1)
            # Set slider position based on selected frame
            self.slider.set(frame)

        def edit_own_annotation(self):
            selected_item = self.annotation_own_tree.selection()
            if len(selected_item) == 1:
                item_values = self.annotation_own_tree.item(selected_item[0], "values")
                frame_selected = item_values[0]
                author_selected = item_values[1]
                note_selected = self.own_full_contents[selected_item[0]]
                self.annotation_edit_window = tk.Toplevel(self.review_toplevel)
                self.annotation_edit_window.title("Annotation Editing")
                # 设置窗口的位置
                x = int(self.screen_width / 2)  # 从左侧边缘的位置（以像素为单位）设置x坐标
                y = int(self.screen_height / 3)  # 从顶部边缘的位置（以像素为单位）设置y坐标
                self.annotation_edit_window.geometry(f"+{x}+{y}")
                self.annotation_edit_label = tk.Label(self.annotation_edit_window, text="Edit your annotation here:")
                self.annotation_edit_label.pack(padx=10, pady=10)
                self.annotation_edit_fa_label = tk.Label(self.annotation_edit_window,
                                                         text=f"Frame: {frame_selected}          Author: {author_selected}")
                self.annotation_edit_fa_label.pack(padx=10, pady=10)
                self.text_edit_entry = tk.Text(self.annotation_edit_window, width=80, height=20)
                self.text_edit_entry.pack(padx=10, pady=10)
                self.ok_edit_button = tk.Button(self.annotation_edit_window, text="Save and Close",
                                                command=self.close_edit_annotation, cursor="hand2")
                self.ok_edit_button.pack(padx=10, pady=10)

                self.text_edit_entry.delete(1.0, tk.END)
                self.text_edit_entry.insert(tk.END, note_selected)

            elif len(selected_item) > 1:
                messagebox.showwarning("Warning", "You can only select one annotation. Please try again.")
            else:
                messagebox.showwarning("Warning", "Please select an annotation.")

        def close_edit_annotation(self):
            selected_item = self.annotation_own_tree.selection()
            item_values = self.annotation_own_tree.item(selected_item[0], "values")
            frame_selected = item_values[0]
            author_selected = item_values[1]
            timestamp_selected = item_values[3]
            text_edited_content = self.text_edit_entry.get("1.0", "end-1c").strip()
            split_text_edited_content = text_edited_content.split("\n")[
                0] if "\n" in text_edited_content else text_edited_content
            # Edit in treeview
            self.annotation_own_tree.set(selected_item, "c3", split_text_edited_content)
            self.own_full_contents[selected_item[0]] = text_edited_content
            self.display_own_text.delete(1.0, tk.END)
            self.display_own_text.insert(tk.END, text_edited_content)
            # Edit in json
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            for i, annotation in enumerate(data["annotation"]):
                if annotation["timestamp"] == timestamp_selected:
                    data["annotation"][i]["text"] = text_edited_content
                    break
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                json.dump(data, file, indent=4)
            # Edit in word
            doc_path_selected = f'{self.master.experiment_path_dynamic}/Annotation/{frame_selected}.docx'
            doc = Document(doc_path_selected)
            # 遍历文档中的所有段落，找到并删除对应的注释
            for para in doc.paragraphs:
                if timestamp_selected in para.text:
                    para.text = f'{timestamp_selected}\n{author_selected}: \n{text_edited_content}\n--------------------------------------------------------------------------------------------------------------------'
            doc.save(doc_path_selected)
            self.annotation_edit_window.destroy()

        def delete_own_annotation(self):
            selected_item = self.annotation_own_tree.selection()
            if len(selected_item) == 1:
                choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete this annotation?")
                if choice:
                    item_values = self.annotation_own_tree.item(selected_item[0], "values")
                    frame_selected = item_values[0]
                    timestamp_selected = item_values[3]
                    # Delete from json
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                        data = json.load(file)
                    for i, annotation in enumerate(data["annotation"]):
                        if annotation["timestamp"] == timestamp_selected:
                            del data["annotation"][i]
                            break
                    with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                        json.dump(data, file, indent=4)
                    # Delete from word
                    doc_path_selected = f'{self.master.experiment_path_dynamic}/Annotation/{frame_selected}.docx'
                    doc = Document(doc_path_selected)
                    # 遍历文档中的所有段落，找到并删除对应的注释
                    for para in doc.paragraphs:
                        if timestamp_selected in para.text:
                            # 找到对应的注释，将其删除
                            p = para._element
                            p.getparent().remove(p)
                    doc.save(doc_path_selected)
                    doc = Document(doc_path_selected)
                    # 判断文档是否为空
                    is_empty = all(not elem.text for elem in doc.element.body)
                    # 如果文档为空，删除文档
                    if is_empty:
                        os.remove(doc_path_selected)
                    # Delete from treeview
                    self.annotation_own_tree.delete(selected_item[0])
                    self.update_sort(self.annotation_own_tree)
                    del self.own_full_contents[selected_item[0]]
                    self.display_own_text.delete(1.0, tk.END)
                else:
                    pass
            elif len(selected_item) > 1:
                choice = messagebox.askyesno("Confirmation", "Are you sure you want to delete these annotations?")
                if choice:
                    # Delete from treeview, json and word
                    for item_id in selected_item:
                        item_values = self.annotation_own_tree.item(item_id, "values")
                        frame_selected = item_values[0]
                        timestamp_selected = item_values[3]
                        # Delete from json
                        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                            data = json.load(file)
                        for i, annotation in enumerate(data["annotation"]):
                            if annotation["timestamp"] == timestamp_selected:
                                del data["annotation"][i]
                                break
                        with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                            json.dump(data, file, indent=4)
                        # Delete from word
                        doc_path_selected = f'{self.master.experiment_path_dynamic}/Annotation/{frame_selected}.docx'
                        doc = Document(doc_path_selected)
                        # 遍历文档中的所有段落，找到并删除对应的注释
                        for para in doc.paragraphs:
                            if timestamp_selected in para.text:
                                # 找到对应的注释，将其删除
                                p = para._element
                                p.getparent().remove(p)
                        doc.save(doc_path_selected)
                        doc = Document(doc_path_selected)
                        # 判断文档是否为空
                        is_empty = all(not elem.text for elem in doc.element.body)
                        # 如果文档为空，删除文档
                        if is_empty:
                            os.remove(doc_path_selected)
                        # Delete from treeview
                        self.annotation_own_tree.delete(item_id)
                        self.update_sort(self.annotation_own_tree)
                        del self.own_full_contents[item_id]
                        self.display_own_text.delete(1.0, tk.END)
                else:
                    pass
            else:
                messagebox.showwarning("Warning", "Please select at least one annotation.")

        def back_to_initial(self):
            # Delete all children of self.right_frame
            for widget in self.right_frame.winfo_children():
                widget.destroy()

                # 在右侧Frame中创建一个内嵌的Frame来装Treeview和Scrollbar
                self.annotation_tree_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
                self.annotation_tree_frame.place(relx=0.5, rely=0.275, relwidth=0.9, relheight=0.5, anchor='center')

                self.annotation_tree = ttk.Treeview(self.annotation_tree_frame, columns=("c1", "c2", "c3", "c4"),
                                                    show='headings')
                self.annotation_tree.column("c1", anchor=tk.CENTER, width=25)
                self.annotation_tree.heading("c1", text="Frame",
                                             command=lambda _col="c1": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c2", anchor=tk.CENTER, width=50)
                self.annotation_tree.heading("c2", text="Author",
                                             command=lambda _col="c2": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c3", anchor=tk.CENTER, width=325)
                self.annotation_tree.heading("c3", text="Notes",
                                             command=lambda _col="c3": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.column("c4", width=0, minwidth=0, stretch='no')
                self.annotation_tree.heading("c4", text="Timestamp",
                                             command=lambda _col="c4": self.sortby(_col, 0, self.annotation_tree))
                self.annotation_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

                self.annotation_tree.bind('<<TreeviewSelect>>', self.on_treeview_select)
                # 创建一个字典用来保存每个 Treeview 项目的完整内容
                self.full_contents = {}

                # Create a vertical scrollbar
                self.annotation_tree_vsb = ttk.Scrollbar(self.annotation_tree_frame, orient="vertical",
                                                         command=self.annotation_tree.yview)
                self.annotation_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)
                # Configure the treeview to use the scrollbar
                self.annotation_tree.configure(yscrollcommand=self.annotation_tree_vsb.set)

                # Add annotation to metadata.json
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)

                for note in data["annotation"]:
                    current_frame = note["frame"]
                    annotation_author = note["author"]
                    text_content = note["text"]
                    timestamp = note["timestamp"]
                    split_text_content = text_content.split("\n")[0] if "\n" in text_content else text_content
                    # Add annotation to treeview
                    item_id = self.annotation_tree.insert('', 'end', values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.full_contents[item_id] = text_content
                self.update_sort(self.annotation_tree)

                # 创建容纳Text的Frame
                self.text_frame = tk.Frame(self.right_frame)  # Create a frame for the Treeview and Scrollbar
                self.text_frame.place(relx=0.495, rely=0.73, relwidth=0.9, relheight=0.36, anchor='center')

                # 创建Text
                self.display_text = tk.Text(self.text_frame, width=40, height=10)
                self.display_text.pack(fill=tk.BOTH, expand=True)

                # 创建manage my own annotations button
                self.manage_own_btn = tk.Button(self.right_frame, text="Manage My Own Annotations",
                                                command=self.manage_own_annotations, cursor="hand2")
                self.manage_own_btn.place(relx=0.5, rely=0.95, anchor='center')

        def on_treeview_select(self, event):
            item_id = self.annotation_tree.focus()
            item = self.annotation_tree.item(item_id)  # Get the item
            frame = item['values'][0]  # The frame value is the first value
            full_content = self.full_contents[item_id]
            self.display_text.delete(1.0, tk.END)
            self.display_text.insert(tk.END, full_content)
            # Set video frame position based on selected frame
            self.cap.set(cv.CAP_PROP_POS_FRAMES, frame - 1)
            # Set slider position based on selected frame
            self.slider.set(frame)

        def make_notes(self):
            current_frame = int(self.slider.get())
            annotation_author = self.master.username.get().strip()
            self.pause_video = True
            self.play_pause_btn.config(text='Play')
            self.annotation_window = tk.Toplevel(self.review_toplevel)
            self.annotation_window.title("Annotation")
            # 设置窗口的位置
            x = int(self.screen_width / 2)  # 从左侧边缘的位置（以像素为单位）设置x坐标
            y = int(self.screen_height / 3)  # 从顶部边缘的位置（以像素为单位）设置y坐标
            self.annotation_window.geometry(f"+{x}+{y}")
            self.annotation_label = tk.Label(self.annotation_window, text="Make your annotations here:")
            self.annotation_label.pack(padx=10, pady=10)
            self.annotation_fa_label = tk.Label(self.annotation_window,
                                                text=f"Frame: {current_frame}          Author: {annotation_author}")
            self.annotation_fa_label.pack(padx=10, pady=10)
            self.text_entry = tk.Text(self.annotation_window, width=80, height=20)
            self.text_entry.pack(padx=10, pady=10)
            self.ok_button = tk.Button(self.annotation_window, text="Save and Close", command=self.close_annotation,
                                       cursor="hand2")
            self.ok_button.pack(padx=10, pady=10)

        def close_annotation(self):
            current_frame = int(self.slider.get())
            annotation_author = self.master.username.get().strip()
            text_content = self.text_entry.get("1.0", "end-1c").strip()
            split_text_content = text_content.split("\n")[0] if "\n" in text_content else text_content
            if text_content == "":
                messagebox.showwarning("Warning", "Please enter valid annotation.")
            else:
                timestamp = datetime.datetime.now().isoformat()
                if self.annotation_tree.winfo_exists():
                    # Add annotation to treeview
                    item_id = self.annotation_tree.insert('', 'end', values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.update_sort(self.annotation_tree)
                    self.full_contents[item_id] = text_content
                elif self.annotation_own_tree.winfo_exists():
                    # Add annotation to treeview
                    item_id = self.annotation_own_tree.insert('', 'end', values=(
                    current_frame, annotation_author, split_text_content, timestamp))
                    self.update_sort(self.annotation_own_tree)
                    self.own_full_contents[item_id] = text_content

                # Add annotation to metadata
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                    data = json.load(file)

                new_annotation = {
                    "frame": current_frame,
                    "author": annotation_author,
                    "text": text_content,
                    "timestamp": timestamp
                }
                if "annotation" not in data:
                    data["annotation"] = [new_annotation]
                else:
                    data["annotation"].append(new_annotation)
                with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'w') as file:
                    json.dump(data, file, indent=4)

                # Add annotation to Annotation file as .docx
                frame_ann_docx_path = f'{self.master.experiment_path_dynamic}/Annotation/{current_frame}.docx'
                if os.path.exists(frame_ann_docx_path):
                    doc = Document(frame_ann_docx_path)
                else:
                    doc = Document()  # 创建一个Document对象
                doc.add_paragraph(
                    f'{timestamp}\n{annotation_author}: \n{text_content}\n--------------------------------------------------------------------------------------------------------------------')  # 在文档中添加一个段落
                doc.save(frame_ann_docx_path)  # 将文档保存为一个.docx文件

                self.pause_video = True
                self.annotation_window.destroy()

        def slider_moved(self, event):
            # Only update frame when video is paused
            if self.pause_video:
                # Change frame position based on slider
                self.cap.set(cv.CAP_PROP_POS_FRAMES, self.slider.get() - 1)
                # Play one frame
                success, frame = self.cap.read()
                if success:
                    height, width, _ = frame.shape
                    aspect_ratio = width / height
                    if width > self.screen_width or height > self.screen_height:
                        if self.screen_width / aspect_ratio <= self.screen_height:
                            new_width = self.screen_width
                            new_height = int(self.screen_width / aspect_ratio)
                        else:
                            new_height = self.screen_height
                            new_width = int(self.screen_height * aspect_ratio)
                        new_height = int(new_height * 3.2 / 5)
                        new_width = int(new_width * 3.2 / 5)
                    else:
                        new_height = int(height * 3.2 / 5)
                        new_width = int(width * 3.2 / 5)
                    # self.master.scale_length = new_width
                    img = cv.cvtColor(frame, cv.COLOR_BGR2RGBA)
                    current_image = PIL.Image.fromarray(img).resize((new_width, new_height), PIL.Image.LANCZOS)
                    enhancer = ImageEnhance.Sharpness(current_image)
                    current_image = enhancer.enhance(2.0)  # Increase sharpness
                    imgtk = PIL.ImageTk.PhotoImage(image=current_image)
                    self.movieLabel.imgtk = imgtk
                    self.movieLabel.config(image=imgtk)
                    self.review_toplevel.update()
            else:
                self.cap.set(cv.CAP_PROP_POS_FRAMES, self.slider.get() - 1)

        def video_player(self):
            while self.master.thread_running:
                read_first_frame = False
                while self.cap.isOpened():
                    if self.pause_video and read_first_frame:
                        continue
                    success, frame = self.cap.read()
                    read_first_frame = True
                    if success == True:
                        height, width, _ = frame.shape
                        aspect_ratio = width / height
                        if width > self.screen_width or height > self.screen_height:
                            if self.screen_width / aspect_ratio <= self.screen_height:
                                new_width = self.screen_width
                                new_height = int(self.screen_width / aspect_ratio)
                            else:
                                new_height = self.screen_height
                                new_width = int(self.screen_height * aspect_ratio)
                            new_height = int(new_height * 3.2 / 5)
                            new_width = int(new_width * 3.2 / 5)
                        else:
                            new_height = int(height * 3.2 / 5)
                            new_width = int(width * 3.2 / 5)
                        # self.master.scale_length = new_width
                        # 使Label框中显示指定帧
                        img = cv.cvtColor(frame, cv.COLOR_BGR2RGBA)
                        current_image = PIL.Image.fromarray(img).resize((new_width, new_height), PIL.Image.LANCZOS)
                        enhancer = ImageEnhance.Sharpness(current_image)
                        current_image = enhancer.enhance(2.0)  # Increase sharpness
                        imgtk = PIL.ImageTk.PhotoImage(image=current_image)
                        try:
                            self.movieLabel.imgtk = imgtk
                            # 在Tkinter的主循环停止后，尝试更新Tkinter组件，这是不允许的。
                            # 在你的案例中，当关闭窗口时，video_player线程仍在运行，并尝试访问已经销毁的Tkinter组件self.movieLabel_1，从而抛出错误。
                            # 你可以在更新Tkinter组件前，检查该组件是否还存在。
                            # 这可以通过使用winfo_exists()函数来实现，这个函数会返回1如果窗口仍然存在，0如果已经被销毁。
                            if self.movieLabel.winfo_exists():
                                self.movieLabel.config(image=imgtk)
                                self.slider.set(self.cap.get(cv.CAP_PROP_POS_FRAMES))
                                self.review_toplevel.update()
                        except:
                            pass
                    else:
                        self.pause_video = True
                        try:
                            if self.movieLabel.winfo_exists():
                                self.play_pause_btn.config(text='Play')
                        except:
                            pass

                self.cap.release()
                time.sleep(1)

        def play_pause(self):  # toggle play/pause status
            self.pause_video = not self.pause_video
            if self.pause_video:
                self.play_pause_btn.config(text='Play')
            else:
                self.play_pause_btn.config(text='Pause')

        def update_process_status(self):
            with open(f'{self.master.experiment_path_dynamic}/metadata.json', 'r') as file:
                data = json.load(file)
            self.master.process_status = data['status']
            self.start_processing_label.config(text=self.master.process_status)


t = Timer(lambda: code_to_test())
t_time = t.timeit(number=1)
print(t_time)
# CPU缓存：当代码多次执行时，其所需的数据可能已经被加载到CPU的缓存中，使得访问速度更快。这尤其在处理数据密集型任务时会更明显。
